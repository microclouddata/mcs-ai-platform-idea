"""
Generate MCS AI Platform Architecture Design Document
Outputs: architecture_design.docx and architecture_design.pdf
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import datetime

TODAY = datetime.date.today().strftime("%B %d, %Y")
VERSION = "1.0"
PROJECT = "MCS AI Platform"

# ─────────────────────────────────────────────────────────────────────────────
# Document content definition
# ─────────────────────────────────────────────────────────────────────────────

SECTIONS = [
    {
        "title": "1. Executive Summary",
        "content": [
            ("para", (
                "The MCS AI Platform is a full-stack, multi-tenant SaaS platform that enables "
                "organisations to create, configure, and deploy AI agents powered by large language "
                "models (LLMs). The platform provides a comprehensive suite of capabilities including "
                "intelligent document search, multi-step workflow automation, scheduled job execution, "
                "prompt template management, usage analytics, and enterprise-grade features such as "
                "organisation management, API key authentication, audit logging, rate limiting, a "
                "public marketplace, and third-party integrations."
            )),
            ("para", (
                "The system is built on a Spring Boot 3.5 backend with MongoDB persistence and a "
                "Next.js 15 frontend. It follows a stateless JWT + API-key authentication model, "
                "async task execution, and a pluggable tool-executor architecture."
            )),
        ]
    },
    {
        "title": "2. System Overview",
        "content": [
            ("heading3", "2.1 Architecture Style"),
            ("para", (
                "The platform employs a monorepo, monolith-first architecture with clear internal "
                "package boundaries that make future microservice extraction straightforward. "
                "The backend exposes a RESTful JSON API; the frontend is a Next.js single-page "
                "application that communicates exclusively through that API."
            )),
            ("heading3", "2.2 Technology Stack"),
            ("table", {
                "headers": ["Layer", "Technology", "Version", "Purpose"],
                "rows": [
                    ["Frontend", "Next.js (App Router)", "15.2.0", "React SSR/SPA, routing, UI"],
                    ["Frontend", "React", "19", "Component model"],
                    ["Frontend", "TypeScript", "5.x", "Type safety"],
                    ["Frontend", "Tailwind CSS", "4.1.0", "Utility-first styling"],
                    ["Backend", "Spring Boot", "3.5.3", "REST API framework"],
                    ["Backend", "Spring Security 6", "6.x", "Authentication & authorisation"],
                    ["Backend", "Spring Data MongoDB", "4.x", "MongoDB ODM"],
                    ["Backend", "Spring @Async / @Scheduled", "—", "Async & scheduled tasks"],
                    ["Backend", "JJWT", "0.12.6", "JWT generation & validation"],
                    ["Backend", "Lombok", "1.18.x", "Boilerplate reduction"],
                    ["AI/ML", "OpenAI API", "Latest", "LLM inference & embeddings"],
                    ["AI/ML", "text-embedding-3-small", "—", "Vector embeddings (1 536-dim)"],
                    ["Database", "MongoDB", "6+", "Primary data store"],
                    ["Runtime", "Java", "21", "JVM runtime"],
                    ["Runtime", "Node.js", "20.9+", "Frontend dev server / build"],
                ],
            }),
            ("heading3", "2.3 Deployment Topology"),
            ("para", (
                "In the current configuration, both the backend (port 8080) and frontend (port 3000) "
                "run on the same host. MongoDB runs locally on the default port 27017. The system is "
                "designed to be containerised and deployed to any cloud provider. CORS is configured "
                "to allow cross-origin requests from the frontend origins."
            )),
        ]
    },
    {
        "title": "3. Backend Architecture",
        "content": [
            ("heading3", "3.1 Package Structure"),
            ("para", "The backend is organised into feature-oriented packages under com.mcs.aiplatform:"),
            ("table", {
                "headers": ["Package", "Responsibility"],
                "rows": [
                    ["auth", "JWT-based registration, login, token generation"],
                    ["user", "User entity, roles (USER / ADMIN)"],
                    ["config", "Security filter chain, CORS, JWT filter, API-key filter, CurrentUser helper"],
                    ["agent", "Agent CRUD, enable/disable, tool & KB bindings"],
                    ["chat", "Chat session & message management, LLM orchestration, tool execution, memory"],
                    ["document", "File upload, text extraction, chunk splitting, embedding, download, delete"],
                    ["embedding", "OpenAI embedding calls, cosine-similarity computation"],
                    ["knowledgebase", "Knowledge base CRUD, document count tracking"],
                    ["tool", "Tool type registry, agent-tool bindings, tool execution facade"],
                    ["tool.executor", "Pluggable executors: KnowledgeSearch, WebSearch, Calculator"],
                    ["memory", "Short-term context (last 10 msgs), LLM-based summarisation"],
                    ["llm", "LLM provider abstraction, OpenAI provider, mock provider"],
                    ["workflow", "Workflow & step builder, async execution engine, execution history"],
                    ["scheduler", "Scheduled job CRUD, cron-style runner (@Scheduled), next-run calculation"],
                    ["template", "Prompt template CRUD, variable extraction ({{var}}), render engine"],
                    ["usage", "Usage log recording, cost estimation, stats aggregation"],
                    ["organization", "Multi-tenant org creation, membership management (OWNER/ADMIN/MEMBER)"],
                    ["billing", "Subscription plans (FREE/PRO/ENTERPRISE), plan limits"],
                    ["apikey", "API key generation (SHA-256 hashed), authentication filter"],
                    ["audit", "Audit log capture (action, resource, IP, user-agent)"],
                    ["ratelimit", "In-memory sliding-window rate limiter per user/plan"],
                    ["marketplace", "Agent/template/workflow marketplace publish, browse, clone"],
                    ["integration", "Third-party integration configs (Slack, Webhook, Email, etc.), trigger service"],
                    ["admin", "Admin endpoints: users, agents, usage logs"],
                    ["common", "BaseEntity, ApiResponse<T>, GlobalExceptionHandler"],
                ],
            }),
            ("heading3", "3.2 Security Architecture"),
            ("para", (
                "Spring Security 6 protects all endpoints through a stateless filter chain. "
                "Two authentication mechanisms run in sequence:"
            )),
            ("bullets", [
                "ApiKeyAuthFilter — checks the X-API-Key request header, hashes the value with SHA-256, "
                "looks up the ApiKey document, and populates the SecurityContext if valid and not expired.",
                "JwtAuthFilter — checks the Authorization: Bearer <token> header, validates the JJWT token, "
                "extracts userId and role claims, and populates the SecurityContext.",
                "If neither header is present the request proceeds unauthenticated and is rejected by "
                "anyRequest().authenticated() unless it targets a permitted path.",
            ]),
            ("para", (
                "Permitted paths (no auth required): /api/auth/**, /api/documents/**, /api/marketplace "
                "(GET browse only), Swagger UI endpoints."
            )),
            ("para", (
                "CORS is handled by a CorsConfigurationSource bean (not WebMvcConfigurer) so that "
                "Spring Security's CORS filter intercepts OPTIONS preflight requests before the "
                "authentication filters run — a critical requirement in Spring Security 6."
            )),
            ("heading3", "3.3 LLM Abstraction"),
            ("para", (
                "The LlmProvider interface decouples the chat and workflow services from any specific "
                "model provider. The OpenAiProvider implements it using Spring's RestClient to call "
                "the OpenAI Chat Completions API. A mock provider returns deterministic responses for "
                "local development without an API key. The active provider is selected via the "
                "app.llm.mock-enabled configuration flag."
            )),
            ("heading3", "3.4 Tool Executor Pattern"),
            ("para", (
                "Tools are implemented as Spring beans that implement the ToolExecutor interface "
                "(supportedType(), execute()). Spring automatically injects all ToolExecutor "
                "implementations into ToolService as a List<ToolExecutor>, enabling zero-config "
                "registration of new tools. Current executors:"
            )),
            ("bullets", [
                "KnowledgeSearchExecutor — generates an embedding for the query, computes cosine "
                "similarity against all document chunks for the agent's knowledge bases, returns "
                "the top-3 relevant chunks. Falls back to keyword search if embeddings are unavailable.",
                "WebSearchExecutor — stub returning a placeholder; ready for integration with "
                "Tavily, SerpAPI, or similar.",
                "CalculatorExecutor — evaluates mathematical expressions using Java's "
                "JavaScript ScriptEngine with sanitisation against dangerous patterns.",
            ]),
            ("heading3", "3.5 Async Workflow Execution"),
            ("para", (
                "@EnableAsync is declared on AiPlatformApplication. WorkflowExecutionService.runAsync() "
                "is annotated with @Async so it executes on Spring's task executor thread pool, "
                "decoupling the HTTP response from execution time. The frontend polls "
                "GET /api/workflows/executions/{id} every 2 seconds until status is COMPLETED or FAILED."
            )),
            ("para", (
                "Each workflow step resolves {{variableName}} placeholders from the execution context "
                "map before execution. Step output is stored back into the context under the step's "
                "outputKey, making it available to subsequent steps."
            )),
            ("heading3", "3.6 Scheduled Jobs"),
            ("para", (
                "WorkflowScheduler is annotated with @Scheduled(fixedRate = 60_000), causing it to "
                "tick every minute. On each tick it queries ScheduledJobRepository.findByEnabledTrue() "
                "and checks whether each job's nextRunAt is in the past. Due jobs trigger "
                "WorkflowService.execute() and update lastRunAt / nextRunAt using ChronoUnit arithmetic "
                "(HOURLY +1h, DAILY +24h, WEEKLY +7d)."
            )),
        ]
    },
    {
        "title": "4. Data Architecture",
        "content": [
            ("heading3", "4.1 MongoDB Collections"),
            ("table", {
                "headers": ["Collection", "Key Fields", "Notes"],
                "rows": [
                    ["users", "email (unique, indexed), passwordHash, name, role", "role: USER | ADMIN"],
                    ["agents", "userId, name, model, systemPrompt, temperature, memoryEnabled, toolsEnabled, knowledgeBaseIds, enabled", "Soft-delete via enabled flag"],
                    ["chat_sessions", "userId, agentId, title, summary", "Summary written by LLM when > 20 messages"],
                    ["chat_messages", "sessionId, role, content, tokens", "role: user | assistant"],
                    ["document_files", "userId, agentId, filename, storagePath, mimeType, chunkCount", "Physical file on local storage"],
                    ["document_chunks", "documentId, chunkIndex, text, embedding (array<double>)", "1 536-dim vector per chunk"],
                    ["knowledge_bases", "userId, name, description, documentCount", "documentCount incremented on upload"],
                    ["agent_tool_bindings", "agentId, toolType, enabled, config", "Compound unique index (agentId, toolType)"],
                    ["workflows", "userId, name, agentId, steps (embedded array), enabled", "steps: id, name, type, inputTemplate, outputKey, config"],
                    ["workflow_executions", "workflowId, userId, status, input, context (map), stepResults (embedded array), finalOutput, error", "status: PENDING|RUNNING|COMPLETED|FAILED"],
                    ["scheduled_jobs", "userId, name, workflowId, scheduleType, enabled, lastRunAt, nextRunAt", "scheduleType: HOURLY|DAILY|WEEKLY"],
                    ["prompt_templates", "userId, name, description, content, variables (array)", "variables extracted from {{name}} patterns"],
                    ["usage_logs", "userId, agentId, model, promptTokens, completionTokens, totalTokens, cost, sessionId", "cost calculated from per-model pricing table"],
                    ["organizations", "name, slug, description, plan, ownerId", "plan: FREE|PRO|ENTERPRISE"],
                    ["org_memberships", "orgId, userId, userEmail, userName, role, invitedBy", "Compound unique index (orgId, userId); role: OWNER|ADMIN|MEMBER"],
                    ["subscriptions", "userId (unique), orgId, plan, status, periodStart, periodEnd, stripeCustomerId", "plan: FREE|PRO|ENTERPRISE"],
                    ["api_keys", "userId, name, keyHash (unique, SHA-256), keyPrefix, scopes, enabled, lastUsedAt, expiresAt", "plaintext key never stored"],
                    ["audit_logs", "userId, orgId, action, resourceType, resourceId, detail, ipAddress, userAgent", "action enum: 20+ action types"],
                    ["marketplace_items", "authorId, type, title, description, category, tags, config, published, downloads, rating", "type: AGENT|TEMPLATE|WORKFLOW"],
                    ["integrations", "userId, orgId, type, name, config (map), enabled, lastTriggeredAt", "type: SLACK|WEBHOOK|EMAIL|WHATSAPP|GOOGLE_DRIVE|NOTION|CONFLUENCE|CRM"],
                ],
            }),
            ("heading3", "4.2 Embedding & Vector Search"),
            ("para", (
                "When a document is uploaded, DocumentService splits the text into chunks (~500 chars "
                "with 100-char overlap), calls EmbeddingService.embed() for each chunk, and stores the "
                "resulting 1 536-dimensional float array alongside the chunk text. At query time, "
                "KnowledgeSearchExecutor embeds the query, iterates over all chunks for the agent's "
                "knowledge bases, computes cosine similarity, and returns the top-3 results above a "
                "0.7 similarity threshold."
            )),
            ("heading3", "4.3 Cost Estimation"),
            ("para", (
                "UsageLogService.calculateCost() uses a static pricing table (per-million-token rates) "
                "to estimate cost at log time. Token counts are estimated as text.length() / 4. "
                "Supported model pricing:"
            )),
            ("table", {
                "headers": ["Model", "Input ($/1M tokens)", "Output ($/1M tokens)"],
                "rows": [
                    ["gpt-4o", "$5.00", "$15.00"],
                    ["gpt-4o-mini", "$0.15", "$0.60"],
                    ["gpt-4.1-mini", "$0.40", "$1.60"],
                    ["gpt-3.5-turbo", "$0.50", "$1.50"],
                ],
            }),
        ]
    },
    {
        "title": "5. Frontend Architecture",
        "content": [
            ("heading3", "5.1 Application Structure"),
            ("para", (
                "The frontend is a Next.js 15 App Router application with all pages under app/. "
                "All interactive pages use the 'use client' directive (client components). "
                "API calls are centralised through the apiFetch<T>() utility in lib/api.ts which "
                "automatically injects the Authorization: Bearer <token> header from localStorage."
            )),
            ("heading3", "5.2 Page Map"),
            ("table", {
                "headers": ["Route", "Page", "Key Features"],
                "rows": [
                    ["/login", "Login", "JWT auth, stores token/userId/name in localStorage"],
                    ["/register", "Register", "User creation"],
                    ["/dashboard", "Dashboard", "Agent cards; disabled agents shown with opacity/grayscale badge"],
                    ["/agents/[id]", "Agent Detail", "Document upload (3-step flow), download, delete; chat link"],
                    ["/agents/[id]/settings", "Agent Settings", "Model, temperature, systemPrompt, memory, tools, KB bindings, danger zone"],
                    ["/knowledge-bases", "Knowledge Bases", "Create/delete KBs"],
                    ["/workflows", "Workflows", "List/create/delete workflows"],
                    ["/workflows/[id]", "Workflow Builder", "Visual step builder, execute with input, real-time polling, execution history"],
                    ["/jobs", "Scheduled Jobs", "Create HOURLY/DAILY/WEEKLY jobs, enable/disable, last/next run times"],
                    ["/templates", "Prompt Templates", "Create with {{variable}} placeholders, render panel with variable inputs"],
                    ["/usage", "Usage & Cost", "Summary cards, token-by-model/agent bar charts, recent log table"],
                    ["/organization", "Organizations", "Create orgs, invite members with roles, manage members"],
                    ["/billing", "Billing", "Plan cards (FREE/PRO/ENTERPRISE), upgrade flow"],
                    ["/api-keys", "API Keys", "Generate keys, one-time plaintext reveal, revoke"],
                    ["/audit", "Audit Logs", "Filterable table with action colour-coding"],
                    ["/marketplace", "Marketplace", "Browse/search/clone published agents, templates, workflows; publish own items"],
                    ["/integrations", "Integrations", "Tile grid for all integration types, dynamic config forms, enable/disable"],
                    ["/admin", "Admin", "Tabbed: Users, Agents, Usage Logs"],
                ],
            }),
            ("heading3", "5.3 Authentication Flow"),
            ("para", (
                "Login stores the JWT in localStorage. NavBar reads localStorage on every route change "
                "(using usePathname as a dependency in useEffect) to reactively show/hide nav links "
                "and the Log out button. Logout removes token, userId, and name from localStorage "
                "and redirects to /login."
            )),
            ("heading3", "5.4 Real-time Polling"),
            ("para", (
                "Workflow execution status is polled via setInterval every 2 seconds on the "
                "workflow detail page. The interval is cleared when execution reaches COMPLETED "
                "or FAILED, and on component unmount via the useEffect cleanup function."
            )),
        ]
    },
    {
        "title": "6. API Reference",
        "content": [
            ("heading3", "6.1 Authentication"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["POST", "/api/auth/register", "None", "Register a new user"],
                    ["POST", "/api/auth/login", "None", "Login, returns JWT"],
                ],
            }),
            ("heading3", "6.2 Agents"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["GET", "/api/agents", "JWT/Key", "List user's agents"],
                    ["POST", "/api/agents", "JWT/Key", "Create agent"],
                    ["GET", "/api/agents/{id}", "JWT/Key", "Get agent by ID"],
                    ["PUT", "/api/agents/{id}", "JWT/Key", "Update agent"],
                    ["PATCH", "/api/agents/{id}/enabled", "JWT/Key", "Toggle agent enabled"],
                    ["DELETE", "/api/agents/{id}", "JWT/Key", "Delete agent"],
                    ["GET", "/api/agents/{id}/tools", "JWT/Key", "Get tool bindings"],
                    ["PUT", "/api/agents/{id}/tools", "JWT/Key", "Update tool bindings"],
                ],
            }),
            ("heading3", "6.3 Chat"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["POST", "/api/chat/{agentId}", "JWT/Key", "Send message, get response"],
                    ["GET", "/api/chat/sessions", "JWT/Key", "List chat sessions"],
                    ["GET", "/api/chat/sessions/{id}/messages", "JWT/Key", "Get messages in session"],
                ],
            }),
            ("heading3", "6.4 Documents"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["POST", "/api/documents/{agentId}/upload", "None*", "Upload document (multipart)"],
                    ["GET", "/api/documents/{agentId}", "None*", "List documents for agent"],
                    ["GET", "/api/documents/{id}/download", "None*", "Download document file"],
                    ["DELETE", "/api/documents/{id}", "None*", "Delete document + chunks"],
                ],
            }),
            ("heading3", "6.5 Workflows"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["GET", "/api/workflows", "JWT/Key", "List workflows"],
                    ["POST", "/api/workflows", "JWT/Key", "Create workflow"],
                    ["GET", "/api/workflows/{id}", "JWT/Key", "Get workflow"],
                    ["PUT", "/api/workflows/{id}", "JWT/Key", "Update workflow"],
                    ["DELETE", "/api/workflows/{id}", "JWT/Key", "Delete workflow"],
                    ["POST", "/api/workflows/{id}/execute", "JWT/Key", "Execute workflow"],
                    ["GET", "/api/workflows/{id}/executions", "JWT/Key", "List executions"],
                    ["GET", "/api/workflows/executions/{execId}", "JWT/Key", "Get execution status"],
                ],
            }),
            ("heading3", "6.6 Platform Features"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["GET", "/api/jobs", "JWT/Key", "List scheduled jobs"],
                    ["POST", "/api/jobs", "JWT/Key", "Create scheduled job"],
                    ["PATCH", "/api/jobs/{id}/enabled", "JWT/Key", "Toggle job enabled"],
                    ["DELETE", "/api/jobs/{id}", "JWT/Key", "Delete job"],
                    ["GET", "/api/templates", "JWT/Key", "List prompt templates"],
                    ["POST", "/api/templates", "JWT/Key", "Create template"],
                    ["POST", "/api/templates/{id}/render", "JWT/Key", "Render template with variables"],
                    ["DELETE", "/api/templates/{id}", "JWT/Key", "Delete template"],
                    ["GET", "/api/usage/stats", "JWT/Key", "Usage statistics"],
                    ["GET", "/api/usage/logs", "JWT/Key", "Usage log entries"],
                    ["GET", "/api/knowledge-bases", "JWT/Key", "List knowledge bases"],
                    ["POST", "/api/knowledge-bases", "JWT/Key", "Create knowledge base"],
                    ["DELETE", "/api/knowledge-bases/{id}", "JWT/Key", "Delete knowledge base"],
                ],
            }),
            ("heading3", "6.7 Enterprise / Phase 4"),
            ("table", {
                "headers": ["Method", "Endpoint", "Auth", "Description"],
                "rows": [
                    ["GET", "/api/organizations", "JWT/Key", "List user's organizations"],
                    ["POST", "/api/organizations", "JWT/Key", "Create organization"],
                    ["GET", "/api/organizations/{id}/members", "JWT/Key", "List org members"],
                    ["POST", "/api/organizations/{id}/members", "JWT/Key", "Invite member"],
                    ["DELETE", "/api/organizations/{id}/members/{userId}", "JWT/Key", "Remove member"],
                    ["GET", "/api/billing/subscription", "JWT/Key", "Get subscription"],
                    ["POST", "/api/billing/subscription/upgrade", "JWT/Key", "Upgrade plan"],
                    ["GET", "/api/api-keys", "JWT/Key", "List API keys"],
                    ["POST", "/api/api-keys", "JWT/Key", "Create API key"],
                    ["DELETE", "/api/api-keys/{id}", "JWT/Key", "Revoke API key"],
                    ["GET", "/api/audit/logs", "JWT/Key", "Get audit logs"],
                    ["GET", "/api/marketplace", "None", "Browse marketplace (public)"],
                    ["POST", "/api/marketplace", "JWT/Key", "Publish item"],
                    ["POST", "/api/marketplace/{id}/clone", "JWT/Key", "Clone marketplace item"],
                    ["GET", "/api/integrations", "JWT/Key", "List integrations"],
                    ["POST", "/api/integrations", "JWT/Key", "Create integration"],
                    ["PATCH", "/api/integrations/{id}/enabled", "JWT/Key", "Toggle integration"],
                    ["DELETE", "/api/integrations/{id}", "JWT/Key", "Delete integration"],
                ],
            }),
        ]
    },
    {
        "title": "7. Security Design",
        "content": [
            ("heading3", "7.1 Authentication"),
            ("bullets", [
                "JWT tokens signed with HS256 using a configurable secret (minimum 32 bytes). "
                "Default expiry: 24 hours (86 400 000 ms).",
                "JWT claims: sub (userId), role (USER|ADMIN), standard exp/iat.",
                "API keys are generated as 68-character random strings prefixed with mcs_. "
                "The SHA-256 hash is stored; the plaintext is returned only once at creation time.",
                "Both JWT and API key auth can be used concurrently. API key filter runs first.",
            ]),
            ("heading3", "7.2 Authorisation"),
            ("bullets", [
                "All endpoints require authentication except: /api/auth/**, /api/documents/**, "
                "/api/marketplace (GET), Swagger UI.",
                "@EnableMethodSecurity is active, enabling @PreAuthorize on individual methods.",
                "Organisation operations enforce membership checks in the service layer.",
                "Agent and document operations scope queries to the authenticated userId.",
            ]),
            ("heading3", "7.3 Rate Limiting"),
            ("para", (
                "RateLimitService maintains a ConcurrentHashMap of per-user sliding daily windows. "
                "Plan limits: FREE = 100 req/day, PRO = 5 000 req/day, ENTERPRISE = 100 000 req/day. "
                "The window resets at midnight UTC."
            )),
            ("heading3", "7.4 Data Protection"),
            ("bullets", [
                "Passwords are hashed with BCrypt (strength 10).",
                "API key plaintext is never persisted — only the SHA-256 hash.",
                "Sensitive integration config fields (tokens, passwords) are stored as-is in MongoDB; "
                "field-level encryption is recommended for production.",
                "CORS allows only configured origins (default: localhost:3000, localhost:3001).",
            ]),
        ]
    },
    {
        "title": "8. Platform Features",
        "content": [
            ("heading3", "8.1 Agent Configuration"),
            ("para", (
                "Each agent is configured with: LLM model, temperature (0.0–2.0), system prompt, "
                "memory enabled flag, tools enabled flag, and a list of knowledge base IDs. "
                "Agents can be soft-disabled (enabled = false); disabled agents appear with a "
                "visual indicator in the dashboard and cannot be chatted with."
            )),
            ("heading3", "8.2 Memory System"),
            ("para", (
                "MemoryService provides two levels of context: short-term context (last 10 messages "
                "injected into the LLM prompt for context continuity) and long-term summarisation "
                "(when a session exceeds 20 messages, the LLM is called to produce a summary that "
                "replaces the message history, keeping the context window manageable)."
            )),
            ("heading3", "8.3 Workflow Engine"),
            ("para", "Supported step types:"),
            ("table", {
                "headers": ["Step Type", "What it does"],
                "rows": [
                    ["LLM_PROMPT", "Sends inputTemplate (after variable resolution) to the agent's LLM with an optional system prompt"],
                    ["KNOWLEDGE_SEARCH", "Runs KnowledgeSearchExecutor against the agent's knowledge bases"],
                    ["WEB_SEARCH", "Runs WebSearchExecutor (currently a stub, ready for real search API integration)"],
                    ["SUMMARIZE", "Asks the LLM to summarise the input"],
                    ["HTTP_REQUEST", "Makes an HTTP POST to a configured URL with the resolved input as body"],
                ],
            }),
            ("heading3", "8.4 Subscription Plans"),
            ("table", {
                "headers": ["Plan", "Max Agents", "Requests/Day", "Max Documents", "Price"],
                "rows": [
                    ["FREE", "5", "100", "50", "$0/month"],
                    ["PRO", "50", "5,000", "1,000", "$29/month"],
                    ["ENTERPRISE", "500", "100,000", "10,000", "$149/month"],
                ],
            }),
            ("heading3", "8.5 Marketplace"),
            ("para", (
                "Users can publish agents, prompt templates, and workflows to the public marketplace. "
                "Published items are discoverable by all platform users without authentication. "
                "Cloning an item increments its download counter. Authors can unpublish their items "
                "at any time."
            )),
            ("heading3", "8.6 Integrations"),
            ("table", {
                "headers": ["Integration", "Mechanism", "Config Fields"],
                "rows": [
                    ["Slack", "Outgoing webhook HTTP POST", "webhookUrl"],
                    ["Webhook (generic)", "Outgoing HTTP POST on workflow events", "webhookUrl, secret"],
                    ["Email", "SMTP (configurable)", "smtpHost, smtpPort, username, password"],
                    ["WhatsApp", "Meta Cloud API", "phoneNumberId, accessToken"],
                    ["Google Drive", "OAuth 2.0", "clientId, clientSecret"],
                    ["Notion", "Notion Integration Token", "integrationToken"],
                    ["Confluence", "Atlassian API Token", "baseUrl, apiToken, email"],
                    ["CRM", "Provider API Key", "provider, apiKey"],
                ],
            }),
        ]
    },
    {
        "title": "9. Development & Operations",
        "content": [
            ("heading3", "9.1 Local Development Setup"),
            ("bullets", [
                "Prerequisites: Java 21, Maven 3.9+, Node.js 20.9+, MongoDB 6+ running locally on port 27017.",
                "Backend: open backend/pom.xml in IntelliJ IDEA, run AiPlatformApplication. "
                "Server starts on http://localhost:8080.",
                "Frontend: cd frontend && npm install && npm run dev. App runs on http://localhost:3000.",
                "Swagger UI: http://localhost:8080/swagger-ui/index.html",
                "Set app.llm.mock-enabled: false and provide OPENAI_API_KEY for real LLM responses.",
            ]),
            ("heading3", "9.2 Configuration Reference"),
            ("table", {
                "headers": ["Property", "Default", "Description"],
                "rows": [
                    ["server.port", "8080", "Backend HTTP port"],
                    ["spring.data.mongodb.uri", "mongodb://localhost:27017/mcs_ai_platform", "MongoDB connection string"],
                    ["app.jwt.secret", "(change this)", "HS256 signing key, must be ≥ 32 bytes"],
                    ["app.jwt.expiration-ms", "86400000", "JWT TTL in milliseconds (24 h)"],
                    ["app.cors.allowed-origins", "localhost:3000, :3001", "Comma-separated allowed CORS origins"],
                    ["app.storage.local-path", "D://file_upload", "File system path for uploaded documents"],
                    ["app.llm.mock-enabled", "false", "true = use mock LLM, false = call OpenAI"],
                    ["app.llm.openai.api-key", "(set via env)", "OpenAI API key"],
                    ["app.llm.openai.model", "gpt-3.5-turbo", "Default LLM model"],
                    ["llm.embedding.model", "text-embedding-3-small", "OpenAI embedding model"],
                ],
            }),
            ("heading3", "9.3 MongoDB Collections Summary"),
            ("para", (
                "Total active collections: 20. All documents extend BaseEntity which provides: "
                "id (MongoDB ObjectId as String), createdAt (Instant, @CreatedDate), "
                "updatedAt (Instant, @LastModifiedDate). Auditing requires @EnableMongoAuditing "
                "on the application class."
            )),
        ]
    },
    {
        "title": "10. Roadmap & Future Considerations",
        "content": [
            ("heading3", "10.1 Near-term Improvements"),
            ("bullets", [
                "Replace local cosine-similarity vector search with MongoDB Atlas Vector Search "
                "for production-scale semantic retrieval.",
                "Add Ollama adapter to LlmProviderFactory for local/self-hosted model support.",
                "Implement streaming chat responses (Server-Sent Events or WebSocket) to improve UX.",
                "Add real web search integration (Tavily or SerpAPI) to WebSearchExecutor.",
                "Add Stripe webhook handling to SubscriptionService for real payment processing.",
                "Implement field-level encryption for sensitive integration config values.",
                "Add email notification system (welcome emails, invitation emails, billing alerts).",
            ]),
            ("heading3", "10.2 Architectural Evolution"),
            ("bullets", [
                "Containerise both services with Docker; add docker-compose for local development.",
                "Extract workflow execution into a dedicated microservice as load increases.",
                "Add a message queue (RabbitMQ / Kafka) between the API and async workers "
                "to improve reliability and enable retry logic.",
                "Introduce Redis for distributed rate limiting and session caching.",
                "Add OpenTelemetry tracing across backend services.",
                "Implement SSO / SAML 2.0 for enterprise organisation members.",
            ]),
            ("heading3", "10.3 Known Limitations (Current Version)"),
            ("bullets", [
                "Token counts are estimated (text.length() / 4) rather than using the "
                "actual counts returned by the OpenAI API.",
                "Rate limiting state is in-memory; restarts reset counters and the system "
                "cannot rate-limit across multiple backend instances.",
                "Document endpoints (/api/documents/**) are permit-all — consider adding "
                "token-based auth for production.",
                "WebSearchExecutor returns a stub response; real search is not yet wired up.",
                "No background retry mechanism for failed workflow steps.",
            ]),
        ]
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# Word document generator
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table_to_doc(doc, data):
    headers = data["headers"]
    rows = data["rows"]
    col_count = len(headers)

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(hdr_cells[i], '2C3E50')

    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        fill = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = str(cell_text)
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_bg(row_cells[c_idx], fill)

    doc.add_paragraph()


def generate_word(output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(PROJECT)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Architecture Design Document")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x34, 0x49, 0x6B)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {VERSION}  ·  {TODAY}").font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Table of contents header
    toc_heading = doc.add_paragraph("Table of Contents")
    toc_heading.style = 'Heading 1'
    for s in SECTIONS:
        toc_para = doc.add_paragraph(s["title"], style='List Number')
        toc_para.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # Content
    for section in SECTIONS:
        heading = doc.add_heading(section["title"], level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)

        for item in section["content"]:
            kind, data = item

            if kind == "para":
                p = doc.add_paragraph(data)
                p.paragraph_format.space_after = Pt(6)

            elif kind == "heading3":
                h = doc.add_heading(data, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x34, 0x49, 0x6B)

            elif kind == "bullets":
                for bullet in data:
                    doc.add_paragraph(bullet, style='List Bullet')

            elif kind == "table":
                add_table_to_doc(doc, data)

    doc.save(output_path)
    print(f"Word document saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF generator
# ─────────────────────────────────────────────────────────────────────────────

NAVY    = colors.HexColor('#1A2B4A')
STEEL   = colors.HexColor('#34496B')
BRAND   = colors.HexColor('#2563EB')
LIGHT   = colors.HexColor('#F8F9FA')
MID     = colors.HexColor('#E2E8F0')
TH_BG   = colors.HexColor('#2C3E50')
TH_FG   = colors.white
ROW_ALT = colors.HexColor('#F1F5F9')

def get_pdf_styles():
    styles = getSampleStyleSheet()
    base = styles['Normal']

    custom = {}

    custom['Title'] = ParagraphStyle('DocTitle',
        parent=base, fontSize=28, leading=36, textColor=NAVY,
        fontName='Helvetica-Bold', spaceAfter=6, alignment=TA_CENTER)

    custom['Subtitle'] = ParagraphStyle('DocSubtitle',
        parent=base, fontSize=16, leading=22, textColor=STEEL,
        fontName='Helvetica', spaceAfter=4, alignment=TA_CENTER)

    custom['Meta'] = ParagraphStyle('DocMeta',
        parent=base, fontSize=10, textColor=colors.grey,
        fontName='Helvetica', spaceAfter=30, alignment=TA_CENTER)

    custom['H1'] = ParagraphStyle('SectionH1',
        parent=base, fontSize=16, leading=22, textColor=NAVY,
        fontName='Helvetica-Bold', spaceBefore=20, spaceAfter=8,
        borderPad=4)

    custom['H2'] = ParagraphStyle('SectionH2',
        parent=base, fontSize=12, leading=18, textColor=STEEL,
        fontName='Helvetica-Bold', spaceBefore=14, spaceAfter=6)

    custom['Body'] = ParagraphStyle('Body',
        parent=base, fontSize=9.5, leading=15, textColor=colors.HexColor('#1E293B'),
        fontName='Helvetica', spaceAfter=8, alignment=TA_JUSTIFY)

    custom['Bullet'] = ParagraphStyle('Bullet',
        parent=base, fontSize=9.5, leading=14, textColor=colors.HexColor('#1E293B'),
        fontName='Helvetica', leftIndent=18, spaceAfter=4,
        bulletIndent=6, bulletFontName='Helvetica')

    custom['TOCTitle'] = ParagraphStyle('TOCTitle',
        parent=base, fontSize=14, textColor=NAVY, fontName='Helvetica-Bold',
        spaceAfter=10)

    custom['TOCItem'] = ParagraphStyle('TOCItem',
        parent=base, fontSize=10, textColor=STEEL, fontName='Helvetica',
        leftIndent=12, spaceAfter=3)

    custom['TH'] = ParagraphStyle('TH',
        parent=base, fontSize=8.5, textColor=TH_FG,
        fontName='Helvetica-Bold', leading=12)

    custom['TD'] = ParagraphStyle('TD',
        parent=base, fontSize=8, textColor=colors.HexColor('#1E293B'),
        fontName='Helvetica', leading=11)

    return custom


def make_pdf_table(headers, rows, styles):
    th_style = styles['TH']
    td_style = styles['TD']

    col_count = len(headers)
    # Distribute width proportionally
    page_w = A4[0] - 4*cm
    col_w = page_w / col_count

    table_data = [[Paragraph(h, th_style) for h in headers]]
    for row in rows:
        table_data.append([Paragraph(str(c), td_style) for c in row])

    col_widths = [col_w] * col_count

    t = Table(table_data, colWidths=col_widths, repeatRows=1)

    row_bg = []
    for i in range(1, len(rows) + 1):
        bg = ROW_ALT if i % 2 == 1 else colors.white
        row_bg.append(('BACKGROUND', (0, i), (-1, i), bg))

    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), TH_BG),
        ('TEXTCOLOR', (0, 0), (-1, 0), TH_FG),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [ROW_ALT, colors.white]),
    ] + row_bg))
    return t


def generate_pdf(output_path):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        title=f"{PROJECT} Architecture Design Document",
        author="MCS Engineering",
    )

    styles = get_pdf_styles()
    story = []

    # ── Title page ──
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(PROJECT, styles['Title']))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph("Architecture Design Document", styles['Subtitle']))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph(f"Version {VERSION}  ·  {TODAY}", styles['Meta']))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BRAND, spaceAfter=20))

    # ── Table of contents ──
    story.append(Paragraph("Table of Contents", styles['TOCTitle']))
    for s in SECTIONS:
        story.append(Paragraph(s["title"], styles['TOCItem']))
    story.append(PageBreak())

    # ── Sections ──
    for section in SECTIONS:
        story.append(Paragraph(section["title"], styles['H1']))
        story.append(HRFlowable(width="100%", thickness=0.5, color=MID, spaceAfter=8))

        for item in section["content"]:
            kind, data = item

            if kind == "para":
                story.append(Paragraph(data, styles['Body']))

            elif kind == "heading3":
                story.append(Paragraph(data, styles['H2']))

            elif kind == "bullets":
                for bullet in data:
                    story.append(Paragraph(f"• {bullet}", styles['Bullet']))
                story.append(Spacer(1, 0.2*cm))

            elif kind == "table":
                tbl = make_pdf_table(data["headers"], data["rows"], styles)
                story.append(KeepTogether([tbl, Spacer(1, 0.4*cm)]))

        story.append(Spacer(1, 0.4*cm))

    # ── Footer function ──
    def add_footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.grey)
        canvas.drawString(2.5*cm, 1.5*cm, f"{PROJECT} — Architecture Design Document — v{VERSION}")
        canvas.drawRightString(A4[0] - 2.5*cm, 1.5*cm, f"Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    print(f"PDF document saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    base = r"D:\Users\billc\Documents\IdeaProjects\mcs-ai-platform-idea"
    word_path = f"{base}\\MCS_AI_Platform_Architecture.docx"
    pdf_path  = f"{base}\\MCS_AI_Platform_Architecture.pdf"

    generate_word(word_path)
    generate_pdf(pdf_path)
    print("\nDone! Both documents are ready.")