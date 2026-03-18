"""
MCS AI Platform — User Guide Generator
Creates realistic dark-theme UI mockup images, then exports the guide as DOCX and PDF.
"""

import os
import math
import textwrap
import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether, Image as RLImage,
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# ─────────────────────────────────────────────────────────────────────────────
# Constants & paths
# ─────────────────────────────────────────────────────────────────────────────
BASE = Path(r"D:\Users\billc\Documents\IdeaProjects\mcs-ai-platform-idea")
IMG_DIR = BASE / "guide_images"
IMG_DIR.mkdir(exist_ok=True)

TODAY    = datetime.date.today().strftime("%B %d, %Y")
VERSION  = "1.0"
PROJECT  = "MCS AI Platform"

# ─────────────────────────────────────────────────────────────────────────────
# Colour palette (mirrors app CSS variables)
# ─────────────────────────────────────────────────────────────────────────────
BG       = (10,  15,  30)     # page background
PANEL    = (20,  28,  50)     # card/panel
PANEL_S  = (26,  36,  62)     # panel-soft
BORDER   = (45,  60,  95)     # border
BRAND    = (37, 99, 235)      # blue accent
BRAND_L  = (96,150,255)       # lighter brand
MUTED    = (100,120,160)      # muted text
TEXT     = (220,230,255)      # primary text
WHITE    = (255,255,255)
RED      = (220,  60,  60)
GREEN    = (34,  197,  94)
YELLOW   = (234, 179,   8)
PURPLE   = (168,  85, 247)

W, H = 1280, 800   # canvas size

# ─────────────────────────────────────────────────────────────────────────────
# Font helpers (fallback to default if system fonts unavailable)
# ─────────────────────────────────────────────────────────────────────────────
def load_font(size=14, bold=False):
    candidates = [
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()

def bold_font(size=14):   return load_font(size, bold=True)
def reg_font(size=14):    return load_font(size, bold=False)

# ─────────────────────────────────────────────────────────────────────────────
# Drawing primitives
# ─────────────────────────────────────────────────────────────────────────────
def new_canvas():
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    return img, d

def rounded_rect(d, x1, y1, x2, y2, r=12, fill=None, outline=None, width=1):
    d.rounded_rectangle([x1, y1, x2, y2], radius=r, fill=fill, outline=outline, width=width)

def draw_navbar(d, active=""):
    """Top navigation bar."""
    d.rectangle([0, 0, W, 52], fill=PANEL)
    d.line([0, 52, W, 52], fill=BORDER, width=1)
    # Logo
    d.text((20, 14), "⬡ MCS AI Platform", font=bold_font(15), fill=BRAND_L)
    # Nav links
    nav_items = ["Dashboard","Knowledge Bases","Workflows","Jobs","Templates",
                 "Usage","Marketplace","Integrations","Admin"]
    x = 240
    for item in nav_items:
        color = WHITE if item == active else MUTED
        d.text((x, 17), item, font=reg_font(12), fill=color)
        if item == active:
            tw = d.textlength(item, font=reg_font(12))
            d.line([x, 50, x + tw, 50], fill=BRAND, width=2)
        x += d.textlength(item, font=reg_font(12)) + 22
    # Log out
    d.text((W - 70, 17), "Log out", font=reg_font(12), fill=MUTED)

def draw_sidebar(d, items, active_idx=0):
    """Left sidebar."""
    d.rectangle([0, 52, 220, H], fill=PANEL)
    d.line([220, 52, 220, H], fill=BORDER, width=1)
    for i, item in enumerate(items):
        y = 80 + i * 40
        if i == active_idx:
            rounded_rect(d, 8, y - 4, 212, y + 28, r=8, fill=BRAND)
            d.text((20, y + 2), item, font=bold_font(13), fill=WHITE)
        else:
            d.text((20, y + 2), item, font=reg_font(13), fill=MUTED)

def draw_card(d, x, y, w, h, r=16, title="", title_size=14, subtitle=""):
    rounded_rect(d, x, y, x+w, y+h, r=r, fill=PANEL, outline=BORDER, width=1)
    if title:
        d.text((x+18, y+16), title, font=bold_font(title_size), fill=TEXT)
    if subtitle:
        d.text((x+18, y+16+title_size+6), subtitle, font=reg_font(11), fill=MUTED)

def draw_input(d, x, y, w, placeholder="", value="", h=42):
    rounded_rect(d, x, y, x+w, y+h, r=10, fill=PANEL_S, outline=BORDER, width=1)
    text = value if value else placeholder
    color = TEXT if value else MUTED
    d.text((x+14, y+12), text, font=reg_font(12), fill=color)

def draw_button(d, x, y, label, w=None, style="primary"):
    tw = d.textlength(label, font=bold_font(13))
    bw = w if w else int(tw + 32)
    if style == "primary":
        rounded_rect(d, x, y, x+bw, y+38, r=10, fill=BRAND)
        d.text((x + (bw - tw)//2, y + 10), label, font=bold_font(13), fill=WHITE)
    elif style == "outline":
        rounded_rect(d, x, y, x+bw, y+38, r=10, outline=BRAND, width=1)
        d.text((x + (bw - tw)//2, y + 10), label, font=bold_font(13), fill=BRAND_L)
    elif style == "danger":
        rounded_rect(d, x, y, x+bw, y+38, r=10, outline=RED, width=1)
        d.text((x + (bw - tw)//2, y + 10), label, font=bold_font(13), fill=RED)
    elif style == "ghost":
        rounded_rect(d, x, y, x+bw, y+38, r=10, outline=BORDER, width=1)
        d.text((x + (bw - tw)//2, y + 10), label, font=bold_font(13), fill=MUTED)
    return bw

def draw_badge(d, x, y, label, color=BRAND, bg=None):
    tw = d.textlength(label, font=reg_font(10))
    bw = int(tw + 16)
    bg_color = bg if bg else tuple(max(0, c - 180) for c in color)
    rounded_rect(d, x, y, x+bw, y+22, r=11, fill=bg_color)
    d.text((x+8, y+4), label, font=reg_font(10), fill=color)
    return bw

def draw_stat_card(d, x, y, w, h, value, label, color=BRAND_L):
    rounded_rect(d, x, y, x+w, y+h, r=16, fill=PANEL, outline=BORDER, width=1)
    d.text((x+20, y+20), value, font=bold_font(26), fill=color)
    d.text((x+20, y+58), label, font=reg_font(12), fill=MUTED)

def draw_progress_bar(d, x, y, total_w, pct, color=BRAND):
    rounded_rect(d, x, y, x+total_w, y+8, r=4, fill=BORDER)
    filled = int(total_w * pct)
    if filled > 4:
        rounded_rect(d, x, y, x+filled, y+8, r=4, fill=color)

def draw_table_row(d, x, y, cols, widths, is_header=False, alt=False):
    row_h = 36
    bg = (30, 42, 70) if is_header else ((22, 32, 55) if alt else PANEL)
    d.rectangle([x, y, x + sum(widths), y + row_h], fill=bg)
    d.line([x, y + row_h, x + sum(widths), y + row_h], fill=BORDER, width=1)
    cx = x
    for i, (col, w) in enumerate(zip(cols, widths)):
        font = bold_font(11) if is_header else reg_font(11)
        color = BRAND_L if is_header else TEXT
        d.text((cx + 12, y + 11), str(col)[:28], font=font, fill=color)
        cx += w

def draw_toggle(d, x, y, on=True):
    bg = GREEN if on else BORDER
    rounded_rect(d, x, y, x+40, y+22, r=11, fill=bg)
    cx = x + 24 if on else x + 8
    d.ellipse([cx-8, y+3, cx+8, y+19], fill=WHITE)

def draw_tag(d, x, y, text, color=PURPLE):
    tw = d.textlength(text, font=reg_font(10))
    bw = int(tw + 16)
    bg = tuple(max(0, c - 180) for c in color)
    rounded_rect(d, x, y, x+bw, y+22, r=11, fill=bg)
    d.text((x+8, y+4), text, font=reg_font(10), fill=color)
    return bw

# ─────────────────────────────────────────────────────────────────────────────
# Screen generators
# ─────────────────────────────────────────────────────────────────────────────

def screen_login():
    img, d = new_canvas()
    # centered card
    cw, ch = 440, 460
    cx = (W - cw) // 2
    cy = (H - ch) // 2 - 20
    rounded_rect(d, cx, cy, cx+cw, cy+ch, r=24, fill=PANEL, outline=BORDER, width=1)
    # Logo
    d.text((cx + cw//2 - 80, cy+30), "⬡ MCS AI Platform", font=bold_font(18), fill=BRAND_L)
    d.text((cx + cw//2 - 40, cy+62), "Sign in", font=bold_font(22), fill=TEXT)
    d.text((cx + cw//2 - 90, cy+94), "Enter your credentials to continue", font=reg_font(12), fill=MUTED)
    # Fields
    d.text((cx+30, cy+136), "Email", font=reg_font(12), fill=MUTED)
    draw_input(d, cx+30, cy+154, cw-60, "you@example.com")
    d.text((cx+30, cy+214), "Password", font=reg_font(12), fill=MUTED)
    draw_input(d, cx+30, cy+232, cw-60, "••••••••••")
    # Button
    rounded_rect(d, cx+30, cy+295, cx+cw-30, cy+340, r=12, fill=BRAND)
    d.text((cx + cw//2 - 28, cy+308), "Sign in", font=bold_font(14), fill=WHITE)
    # Divider
    d.line([cx+30, cy+360, cx+cw-30, cy+360], fill=BORDER, width=1)
    d.text((cx + cw//2 - 85, cy+375), "Don't have an account?", font=reg_font(12), fill=MUTED)
    d.text((cx + cw//2 + 42, cy+375), "Register", font=bold_font(12), fill=BRAND_L)
    return img


def screen_dashboard():
    img, d = new_canvas()
    draw_navbar(d, "Dashboard")
    # Page title
    d.text((30, 70), "Dashboard", font=bold_font(26), fill=TEXT)
    d.text((30, 104), "Your AI agents", font=reg_font(13), fill=MUTED)
    draw_button(d, W-150, 72, "+ New Agent")

    agents = [
        ("Customer Support Bot", "gpt-4o", True,  "Handles customer queries with empathy"),
        ("Code Reviewer",        "gpt-4o-mini", True, "Reviews pull requests and code quality"),
        ("Data Analyst",         "gpt-3.5-turbo", False, "Analyses CSV and database query results"),
        ("Content Writer",       "gpt-4o",     True,  "Generates blog posts and marketing copy"),
        ("SEO Assistant",        "gpt-4o-mini", True, "Keyword research and on-page optimisation"),
        ("Email Drafter",        "gpt-3.5-turbo", True, "Drafts professional email responses"),
    ]

    cols = 3
    card_w, card_h = 370, 150
    gap = 20
    start_x, start_y = 30, 140

    for i, (name, model, enabled, desc) in enumerate(agents):
        row, col = i // cols, i % cols
        x = start_x + col * (card_w + gap)
        y = start_y + row * (card_h + gap)
        alpha = 1 if enabled else 0.4

        rounded_rect(d, x, y, x+card_w, y+card_h, r=16,
                     fill=PANEL if enabled else (15, 22, 40), outline=BORDER, width=1)
        # Status dot
        dot_color = GREEN if enabled else MUTED
        d.ellipse([x+14, y+18, x+24, y+28], fill=dot_color)
        d.text((x+32, y+14), name, font=bold_font(14),
               fill=TEXT if enabled else MUTED)
        if not enabled:
            draw_badge(d, x+card_w-88, y+12, "Disabled", color=(180,180,180), bg=(40,50,70))
        d.text((x+16, y+42), desc, font=reg_font(11), fill=MUTED)
        d.text((x+16, y+66), f"Model: {model}", font=reg_font(11), fill=MUTED)

        bw1 = draw_button(d, x+16, y+100, "Chat", style="primary")
        draw_button(d, x+16+bw1+10, y+100, "Settings", style="outline")
    return img


def screen_agent_detail():
    img, d = new_canvas()
    draw_navbar(d, "Dashboard")

    # Header
    d.text((30, 68), "Customer Support Bot", font=bold_font(22), fill=TEXT)
    draw_badge(d, 310, 74, "gpt-4o", color=BRAND_L)
    draw_badge(d, 380, 74, "Enabled", color=GREEN)
    draw_button(d, W-270, 62, "Settings", style="outline")
    draw_button(d, W-160, 62, "Start Chat")

    # Two column layout
    # Left: Documents
    lw = 560
    draw_card(d, 30, 110, lw, 380, title="Knowledge Documents")
    docs = [
        ("product_manual.pdf",      "PDF", "2.3 MB", "2025-03-01"),
        ("faq_database.txt",        "TXT", "145 KB", "2025-03-05"),
        ("support_playbook.md",     "MD",  "88 KB",  "2025-03-10"),
        ("pricing_guide_2025.pdf",  "PDF", "1.1 MB", "2025-03-12"),
    ]
    draw_table_row(d, 46, 148, ["Filename","Type","Size","Uploaded"], [220,60,80,120], is_header=True)
    for i, (fn, ft, sz, dt) in enumerate(docs):
        draw_table_row(d, 46, 184 + i*36, [fn, ft, sz, dt], [220,60,80,120], alt=i%2==0)

    # Upload area
    rounded_rect(d, 46, 334, 46+lw-32, 460, r=12, outline=BORDER, width=1)
    d.text((150, 370), "⬆  Drag & drop or click to upload", font=reg_font(13), fill=MUTED)
    d.text((190, 398), "PDF, TXT, Markdown — max 20 MB", font=reg_font(11), fill=MUTED)
    draw_button(d, 46, 468, "Upload Document", style="outline")

    # Right: Chat preview
    rw = 600
    rx = 30 + lw + 20
    draw_card(d, rx, 110, rw, 380, title="Recent Chat")

    messages = [
        ("user", "What is your return policy?"),
        ("bot",  "Our return policy allows returns within 30 days of purchase. Items must be in original condition with receipt."),
        ("user", "How do I initiate a return?"),
        ("bot",  "Visit our website > Orders > Select order > Click 'Return'. You'll receive a prepaid label within 24h."),
    ]
    my = 150
    for role, msg in messages:
        is_user = role == "user"
        bw = rw - 80
        lines = textwrap.wrap(msg, 60)
        lh = len(lines) * 18 + 20
        mx = rx + (rw - bw - 10) if is_user else rx + 24
        fill = BRAND if is_user else PANEL_S
        rounded_rect(d, mx, my, mx+bw, my+lh, r=10, fill=fill)
        for li, line in enumerate(lines):
            d.text((mx+12, my+10+li*18), line, font=reg_font(11), fill=WHITE)
        my += lh + 10

    # Input bar
    draw_input(d, rx+16, 462, rw-90, "Type your message…")
    draw_button(d, rx+rw-70, 462, "Send", w=60)
    return img


def screen_agent_settings():
    img, d = new_canvas()
    draw_navbar(d, "Dashboard")
    d.text((30, 68), "Agent Settings — Customer Support Bot", font=bold_font(20), fill=TEXT)

    # Left column — config
    lw = 550
    draw_card(d, 30, 108, lw, 460, title="Configuration")

    fields = [
        ("Agent Name",   "Customer Support Bot"),
        ("Description",  "Handles all inbound support queries"),
        ("System Prompt","You are a helpful support agent. Be concise and empathetic."),
        ("Model",        "gpt-4o"),
        ("Temperature",  "0.7"),
    ]
    y = 150
    for label, val in fields:
        d.text((50, y), label, font=reg_font(11), fill=MUTED)
        h = 60 if label == "System Prompt" else 42
        draw_input(d, 50, y+18, lw-40, value=val, h=h)
        y += h + 28

    # Toggles
    d.text((50, y+4), "Memory Enabled", font=reg_font(12), fill=TEXT)
    draw_toggle(d, lw-30, y, on=True)
    y += 38
    d.text((50, y+4), "Tools Enabled", font=reg_font(12), fill=TEXT)
    draw_toggle(d, lw-30, y, on=True)

    # Right column — tools & KB
    rw = 600
    rx = 30 + lw + 20

    draw_card(d, rx, 108, rw, 220, title="Tools")
    tools = [("Knowledge Search", True), ("Web Search", False), ("Calculator", True)]
    for i, (tool, enabled) in enumerate(tools):
        ty = 156 + i*54
        rounded_rect(d, rx+20, ty, rx+rw-20, ty+42, r=10, fill=PANEL_S, outline=BORDER, width=1)
        d.text((rx+36, ty+12), tool, font=reg_font(13), fill=TEXT)
        draw_toggle(d, rx+rw-60, ty+10, on=enabled)

    draw_card(d, rx, 342, rw, 220, title="Knowledge Bases")
    kbs = [("Product Documentation KB", True), ("FAQ Database KB", True), ("Legal Docs KB", False)]
    for i, (kb, linked) in enumerate(kbs):
        ky = 390 + i*52
        rounded_rect(d, rx+20, ky, rx+rw-20, ky+42, r=10, fill=PANEL_S, outline=BORDER, width=1)
        d.text((rx+36, ky+12), kb, font=reg_font(12), fill=TEXT)
        draw_toggle(d, rx+rw-60, ky+10, on=linked)

    draw_button(d, 50, 576, "Save Changes")
    draw_button(d, 50+160, 576, "Delete Agent", style="danger")
    return img


def screen_knowledge_bases():
    img, d = new_canvas()
    draw_navbar(d, "Knowledge Bases")
    d.text((30, 68), "Knowledge Bases", font=bold_font(24), fill=TEXT)

    # Create KB form
    draw_card(d, 30, 108, W-60, 120, title="Create Knowledge Base")
    draw_input(d, 50, 148, 380, "Name (e.g. Product Documentation)")
    draw_input(d, 450, 148, 500, "Description (optional)")
    draw_button(d, 966, 148, "Create")

    # KB list
    kbs = [
        ("Product Documentation KB", "All product manuals and specs", 12, "2025-02-15"),
        ("FAQ Database KB",          "Frequently asked questions",    5,  "2025-02-20"),
        ("Legal Docs KB",            "Terms, privacy policy, GDPR",  8,  "2025-03-01"),
        ("Support Playbooks KB",     "Agent escalation procedures",   3,  "2025-03-08"),
    ]
    draw_card(d, 30, 248, W-60, 360, title="Your Knowledge Bases")
    draw_table_row(d, 46, 292, ["Name","Description","Documents","Created",""], [300,380,120,160,100], is_header=True)
    for i, (name, desc, docs, created) in enumerate(kbs):
        y = 328 + i*52
        draw_table_row(d, 46, y, [name, desc, str(docs), created, ""], [300,380,120,160,100], alt=i%2==0)
        draw_button(d, 46+1060+8, y+8, "Delete", style="danger")
    return img


def screen_workflows():
    img, d = new_canvas()
    draw_navbar(d, "Workflows")
    d.text((30, 68), "Workflows", font=bold_font(24), fill=TEXT)

    # Create section
    draw_card(d, 30, 108, W-60, 110)
    d.text((50, 120), "Create New Workflow", font=bold_font(15), fill=TEXT)
    draw_input(d, 50, 150, 500, "Workflow name")
    draw_input(d, 566, 150, 500, "Description (optional)")
    draw_button(d, 1082, 150, "Create")

    # Workflow cards
    workflows = [
        ("Customer Onboarding Flow",  "3 steps", "Searches KB → Generates welcome email → Sends via webhook"),
        ("Daily Report Generator",    "4 steps", "Queries data → Summarises → Formats → Sends to Slack"),
        ("Support Ticket Triage",     "2 steps", "Classifies intent → Routes to appropriate team"),
        ("SEO Content Pipeline",      "5 steps", "Research → Outline → Draft → Review → Publish"),
    ]

    draw_card(d, 30, 238, W-60, 460, title="Your Workflows")
    card_w = (W - 80) // 2 - 10
    for i, (name, steps, desc) in enumerate(workflows):
        row, col = i // 2, i % 2
        cx = 46 + col*(card_w+20)
        cy = 282 + row*190
        rounded_rect(d, cx, cy, cx+card_w, cy+170, r=14, fill=PANEL_S, outline=BORDER, width=1)
        d.text((cx+18, cy+18), name, font=bold_font(14), fill=TEXT)
        draw_badge(d, cx+18, cy+46, steps, color=BRAND_L)
        # wrap desc
        for li, line in enumerate(textwrap.wrap(desc, 55)):
            d.text((cx+18, cy+76+li*18), line, font=reg_font(11), fill=MUTED)
        draw_button(d, cx+18, cy+128, "Open", style="outline")
        draw_button(d, cx+18+88, cy+128, "Delete", style="danger")
    return img


def screen_workflow_builder():
    img, d = new_canvas()
    draw_navbar(d, "Workflows")
    d.text((30, 68), "Customer Onboarding Flow", font=bold_font(20), fill=TEXT)
    draw_button(d, W-200, 62, "Save", style="primary")
    draw_button(d, W-280, 62, "Run", style="outline")

    # Steps panel
    sw = 680
    draw_card(d, 30, 108, sw, 470, title="Steps")
    draw_button(d, sw-76, 120, "+ Add step", style="outline")

    steps = [
        ("1", "Search Knowledge Base", "KNOWLEDGE_SEARCH", "{{input}}",        "search_result"),
        ("2", "Generate Email",        "LLM_PROMPT",       "{{search_result}}", "email_draft"),
        ("3", "Send via Webhook",      "HTTP_REQUEST",      "{{email_draft}}",  "webhook_resp"),
    ]
    for i, (num, name, stype, inp, outp) in enumerate(steps):
        sy = 158 + i*130
        rounded_rect(d, 50, sy, 50+sw-40, sy+112, r=12, fill=PANEL_S, outline=BORDER, width=1)
        # Step number circle
        d.ellipse([58, sy+8, 84, sy+34], fill=BRAND)
        d.text((65, sy+13), num, font=bold_font(13), fill=WHITE)
        d.text((94, sy+12), name, font=bold_font(13), fill=TEXT)
        type_colors = {"KNOWLEDGE_SEARCH": BRAND_L, "LLM_PROMPT": PURPLE, "HTTP_REQUEST": YELLOW}
        draw_badge(d, 94, sy+36, stype, color=type_colors.get(stype, MUTED))
        d.text((94, sy+68), f"Input:  {inp}", font=reg_font(11), fill=MUTED)
        d.text((94, sy+86), f"Output: {outp}", font=reg_font(11), fill=MUTED)
        draw_button(d, 50+sw-110, sy+8, "✕", style="ghost")

    # Execution panel
    rx = 30 + sw + 20
    rw = W - rx - 30
    draw_card(d, rx, 108, rw, 220, title="Execute")
    d.text((rx+18, 150), "Initial input (available as {{input}})", font=reg_font(12), fill=MUTED)
    rounded_rect(d, rx+18, 172, rx+rw-18, 248, r=12, fill=PANEL_S, outline=BORDER, width=1)
    d.text((rx+32, 185), "Explain our product features and pricing to a", font=reg_font(11), fill=MUTED)
    d.text((rx+32, 203), "new enterprise customer from the healthcare sector.", font=reg_font(11), fill=MUTED)
    draw_button(d, rx+18, 258, "Run Workflow")

    draw_card(d, rx, 344, rw, 234, title="Last Execution")
    exsteps = [("Search Knowledge Base","✓ SUCCESS","210ms"), ("Generate Email","✓ SUCCESS","1840ms"), ("Send via Webhook","✓ SUCCESS","320ms")]
    for i, (sn, st, dur) in enumerate(exsteps):
        ey = 392 + i*50
        sc = GREEN if "SUCCESS" in st else RED
        rounded_rect(d, rx+18, ey, rx+rw-18, ey+38, r=8, fill=PANEL_S)
        d.ellipse([rx+32, ey+12, rx+44, ey+24], fill=GREEN)
        d.text((rx+52, ey+10), sn, font=reg_font(12), fill=TEXT)
        d.text((rx+rw-120, ey+10), dur, font=reg_font(11), fill=MUTED)
    return img


def screen_scheduled_jobs():
    img, d = new_canvas()
    draw_navbar(d, "Jobs")
    d.text((30, 68), "Scheduled Jobs", font=bold_font(24), fill=TEXT)

    draw_card(d, 30, 108, W-60, 110)
    d.text((50, 120), "Schedule a New Job", font=bold_font(15), fill=TEXT)
    draw_input(d, 50, 150, 280, "Job name")
    # workflow select
    rounded_rect(d, 346, 150, 660, 192, r=10, fill=PANEL_S, outline=BORDER, width=1)
    d.text((360, 162), "Daily Report Generator", font=reg_font(12), fill=TEXT)
    d.text((648, 162), "▾", font=reg_font(12), fill=MUTED)
    # schedule select
    rounded_rect(d, 676, 150, 860, 192, r=10, fill=PANEL_S, outline=BORDER, width=1)
    d.text((692, 162), "Daily", font=reg_font(12), fill=TEXT)
    d.text((842, 162), "▾", font=reg_font(12), fill=MUTED)
    draw_button(d, 876, 150, "Schedule")

    jobs = [
        ("Daily Report",        "Customer Onboarding Flow", "DAILY",  True,  "2025-03-17 08:00", "2025-03-18 08:00"),
        ("Weekly Digest",       "SEO Content Pipeline",    "WEEKLY", True,  "2025-03-10 09:00", "2025-03-17 09:00"),
        ("Hourly Sync",         "Support Ticket Triage",   "HOURLY", False, "2025-03-17 14:00", "2025-03-17 15:00"),
        ("Monthly Summary",     "Daily Report Generator",  "WEEKLY", True,  "2025-03-01 06:00", "2025-03-08 06:00"),
    ]
    draw_card(d, 30, 238, W-60, 440, title="Active Jobs")
    for i, (name, wf, stype, enabled, last, nxt) in enumerate(jobs):
        jy = 284 + i*96
        rounded_rect(d, 46, jy, W-46, jy+80, r=12, fill=PANEL_S, outline=BORDER, width=1)
        dot = GREEN if enabled else MUTED
        d.ellipse([62, jy+16, 74, jy+28], fill=dot)
        d.text((82, jy+12), name, font=bold_font(14), fill=TEXT)
        draw_badge(d, 82, jy+36, stype, color=BRAND_L)
        d.text((180, jy+38), wf, font=reg_font(11), fill=MUTED)
        d.text((82, jy+60), f"Last: {last}  ·  Next: {nxt}", font=reg_font(11), fill=MUTED)
        lbl = "Disable" if enabled else "Enable"
        col = YELLOW if enabled else GREEN
        draw_button(d, W-200, jy+22, lbl, style="outline")
        draw_button(d, W-100, jy+22, "Delete", style="danger")
    return img


def screen_templates():
    img, d = new_canvas()
    draw_navbar(d, "Templates")

    lw = 640
    draw_card(d, 30, 20, lw, 760, title="Prompt Templates")

    d.text((50, 62), "Template Name", font=reg_font(12), fill=MUTED)
    draw_input(d, 50, 80, lw-40, "Email Campaign Template")
    d.text((50, 132), "Description", font=reg_font(12), fill=MUTED)
    draw_input(d, 50, 150, lw-40, "Generates personalised email campaigns")
    d.text((50, 202), "Template Content", font=reg_font(12), fill=MUTED)
    rounded_rect(d, 50, 220, lw-10, 340, r=10, fill=PANEL_S, outline=BORDER, width=1)
    lines = [
        "Write a {{tone}} email about {{topic}} for",
        "{{audience}}. The email should highlight",
        "{{key_benefit}} and include a call to action",
        "for {{cta}}. Keep it under {{word_count}} words.",
    ]
    for li, line in enumerate(lines):
        d.text((64, 234 + li*22), line, font=reg_font(11), fill=TEXT)

    d.text((50, 352), "Detected variables:", font=reg_font(11), fill=MUTED)
    vars_ = ["tone","topic","audience","key_benefit","cta","word_count"]
    vx = 50
    for v in vars_:
        vw = draw_tag(d, vx, 372, v, color=BRAND_L)
        vx += vw + 8
        if vx > lw - 60:
            break

    draw_button(d, 50, 408, "Create Template")

    # Template list
    tmpls = [
        ("Cold Outreach Email", ["tone","name","company","pain_point"]),
        ("Blog Post Outline",   ["topic","audience","keywords","length"]),
        ("Support Response",    ["issue_type","resolution","next_steps"]),
    ]
    d.text((50, 464), "Your Templates", font=bold_font(14), fill=TEXT)
    for i, (tname, tvars) in enumerate(tmpls):
        ty = 492 + i*84
        rounded_rect(d, 50, ty, lw-10, ty+72, r=10, fill=PANEL_S, outline=BORDER, width=1)
        d.text((66, ty+14), tname, font=bold_font(12), fill=TEXT)
        draw_button(d, lw-94, ty+14, "Delete", style="danger")
        vx2 = 66
        for v in tvars[:4]:
            vw = draw_tag(d, vx2, ty+42, v)
            vx2 += vw + 6

    # Right: Render panel
    rw = W - lw - 70
    rx = lw + 50
    draw_card(d, rx, 20, rw, 760, title="Render: Cold Outreach Email")
    d.text((rx+18, 64), "Fill in the variables to preview the output:", font=reg_font(11), fill=MUTED)
    var_fields = [("tone","Professional"),("name","Sarah Johnson"),("company","HealthTech Corp"),("pain_point","Manual data entry")]
    vy = 86
    for label, val in var_fields:
        d.text((rx+18, vy), label, font=reg_font(11), fill=MUTED)
        draw_input(d, rx+18, vy+18, rw-36, value=val, h=36)
        vy += 62
    draw_button(d, rx+18, vy+8, "Render Template")
    rounded_rect(d, rx+18, vy+56, rx+rw-18, vy+260, r=12, fill=(10,18,38), outline=BORDER, width=1)
    result_lines = [
        "Subject: Transform HealthTech Corp's",
        "Workflow — Let's Talk",
        "",
        "Dear Sarah Johnson,",
        "",
        "I noticed that HealthTech Corp might be",
        "struggling with manual data entry...",
        "",
        "Our platform has helped 200+ companies",
        "eliminate this pain point completely.",
    ]
    for li, line in enumerate(result_lines):
        d.text((rx+30, vy+68+li*20), line, font=reg_font(11), fill=TEXT)
    return img


def screen_usage():
    img, d = new_canvas()
    draw_navbar(d, "Usage")
    d.text((30, 68), "Usage & Cost", font=bold_font(24), fill=TEXT)

    # Stat cards
    stats = [
        ("1,284",   "Total requests",  BRAND_L),
        ("2,847,320","Total tokens",   PURPLE),
        ("$12.34",   "Estimated cost", GREEN),
    ]
    cw = (W - 80) // 3 - 10
    for i, (val, lbl, col) in enumerate(stats):
        draw_stat_card(d, 30 + i*(cw+15), 110, cw, 95, val, lbl, col)

    # Charts row
    mw = (W - 80) // 2 - 10
    draw_card(d, 30, 222, mw, 240, title="Tokens by Model")
    models = [("gpt-4o", 0.72, BRAND), ("gpt-4o-mini", 0.20, PURPLE), ("gpt-3.5-turbo", 0.08, YELLOW)]
    for i, (model, pct, col) in enumerate(models):
        my = 270 + i*56
        d.text((50, my), model, font=reg_font(12), fill=TEXT)
        d.text((mw-20, my), f"{int(pct*100)}%", font=reg_font(11), fill=MUTED)
        draw_progress_bar(d, 50, my+22, mw-70, pct, col)

    draw_card(d, 30+mw+20, 222, mw, 240, title="Tokens by Agent")
    agents = [("Customer Support Bot",0.55,BRAND),("Code Reviewer",0.28,PURPLE),("Content Writer",0.17,GREEN)]
    for i, (agent, pct, col) in enumerate(agents):
        ay = 270 + i*56
        ax = 30+mw+20
        d.text((ax+20, ay), agent, font=reg_font(12), fill=TEXT)
        d.text((ax+mw-40, ay), f"{int(pct*100)}%", font=reg_font(11), fill=MUTED)
        draw_progress_bar(d, ax+20, ay+22, mw-70, pct, col)

    # Log table
    draw_card(d, 30, 478, W-60, 280, title="Recent Logs")
    headers = ["Agent","Model","Tokens","Cost","Date"]
    widths = [250, 180, 120, 120, 220]
    draw_table_row(d, 46, 520, headers, widths, is_header=True)
    logs = [
        ("Customer Support Bot","gpt-4o",     "3,240","$0.00081","2025-03-18 14:22"),
        ("Code Reviewer",       "gpt-4o-mini","1,820","$0.00009","2025-03-18 13:58"),
        ("Content Writer",      "gpt-4o",     "5,640","$0.00141","2025-03-18 13:30"),
        ("Customer Support Bot","gpt-4o",     "2,100","$0.00052","2025-03-18 12:45"),
        ("Code Reviewer",       "gpt-4o-mini","980",  "$0.00005","2025-03-18 12:10"),
    ]
    for i, row in enumerate(logs):
        y = 556 + i*36
        draw_table_row(d, 46, y, list(row), widths, alt=i%2==0)
    return img


def screen_organization():
    img, d = new_canvas()
    draw_navbar(d, "Organization")
    d.text((30, 68), "Organizations", font=bold_font(24), fill=TEXT)

    # Sidebar orgs
    sw = 240
    draw_card(d, 30, 108, sw, 580, title="Your Orgs")
    orgs = [("Acme Corp","PRO",True),("Personal","FREE",False)]
    for i, (name, plan, active) in enumerate(orgs):
        oy = 154 + i*72
        fill = PANEL_S if not active else (20,50,120)
        brd = BRAND if active else BORDER
        rounded_rect(d, 46, oy, 46+sw-32, oy+56, r=10, fill=fill, outline=brd, width=1 if not active else 2)
        d.text((60, oy+10), name, font=bold_font(13), fill=TEXT)
        draw_badge(d, 60, oy+32, plan, color=BRAND_L if plan=="PRO" else MUTED)
    draw_button(d, 46, 680, "+ Create Org", style="outline")

    # Main panel
    rx = 30 + sw + 20
    rw = W - rx - 30
    draw_card(d, rx, 108, rw, 200, title="Acme Corp")
    draw_badge(d, rx+rw-100, 122, "PRO", color=BRAND_L, bg=(10,30,90))
    d.text((rx+20, 158), "Invite a new member:", font=bold_font(13), fill=TEXT)
    draw_input(d, rx+20, 182, rw-200, "team@acmecorp.com")
    rounded_rect(d, rx+rw-170, 182, rx+rw-80, 224, r=10, fill=PANEL_S, outline=BORDER, width=1)
    d.text((rx+rw-158, 194), "Member", font=reg_font(12), fill=TEXT)
    draw_button(d, rx+rw-72, 182, "Invite", w=60)

    # Members table
    draw_card(d, rx, 322, rw, 350, title="Members (4)")
    members = [
        ("Bill Chen",       "bill@acmecorp.com",   "OWNER",  None),
        ("Sarah Johnson",   "sarah@acmecorp.com",  "ADMIN",  "Remove"),
        ("Mike Davis",      "mike@acmecorp.com",   "MEMBER", "Remove"),
        ("Lisa Wang",       "lisa@acmecorp.com",   "MEMBER", "Remove"),
    ]
    role_colors = {"OWNER": YELLOW, "ADMIN": BRAND_L, "MEMBER": MUTED}
    for i, (name, email, role, action) in enumerate(members):
        my = 370 + i*66
        rounded_rect(d, rx+18, my, rx+rw-18, my+52, r=10, fill=PANEL_S, outline=BORDER, width=1)
        d.ellipse([rx+32, my+12, rx+52, my+32], fill=BRAND)
        initial = name[0].upper()
        d.text((rx+36, my+15), initial, font=bold_font(12), fill=WHITE)
        d.text((rx+60, my+12), name, font=bold_font(13), fill=TEXT)
        d.text((rx+60, my+30), email, font=reg_font(11), fill=MUTED)
        rc = role_colors.get(role, MUTED)
        draw_badge(d, rx+rw-160, my+16, role, color=rc)
        if action:
            draw_button(d, rx+rw-80, my+12, action, style="danger")
    return img


def screen_billing():
    img, d = new_canvas()
    draw_navbar(d, "Billing")
    d.text((30, 68), "Billing & Subscription", font=bold_font(24), fill=TEXT)

    # Current plan banner
    draw_card(d, 30, 108, W-60, 90)
    d.text((50, 126), "Current Plan:", font=reg_font(13), fill=MUTED)
    d.text((160, 122), "PRO", font=bold_font(22), fill=BRAND_L)
    d.text((222, 126), "· Active", font=reg_font(13), fill=GREEN)
    d.text((50, 160), "Renews April 18, 2025  ·  5,000 requests/day remaining today", font=reg_font(12), fill=MUTED)

    # Plan cards
    plans = [
        ("Free",       "$0/mo",   MUTED,   ["5 agents","100 req/day","50 documents","Community support"],           False),
        ("Pro",        "$29/mo",  BRAND,   ["50 agents","5,000 req/day","1,000 documents","Priority support","API access"], True),
        ("Enterprise", "$149/mo", PURPLE,  ["500 agents","100,000 req/day","10,000 documents","Dedicated support","SSO / SAML","SLA guarantee"], False),
    ]
    pw = (W - 80) // 3 - 10
    for i, (name, price, col, feats, current) in enumerate(plans):
        px = 30 + i*(pw+15)
        py = 214
        brd = col if current else BORDER
        brd_w = 2 if current else 1
        rounded_rect(d, px, py, px+pw, py+440, r=18, fill=PANEL, outline=brd, width=brd_w)
        d.text((px+24, py+24), name, font=bold_font(18), fill=col)
        d.text((px+24, py+54), price, font=bold_font(22), fill=TEXT)
        if current:
            draw_badge(d, px+pw-120, py+24, "Current Plan", color=BRAND_L)
        d.line([px+16, py+90, px+pw-16, py+90], fill=BORDER, width=1)
        for fi, feat in enumerate(feats):
            fy = py+106 + fi*38
            d.text((px+46, fy+2), feat, font=reg_font(12), fill=MUTED)
            d.text((px+26, fy), "✓", font=bold_font(12), fill=GREEN)
        btn_style = "ghost" if current else "primary"
        btn_label = "Current Plan" if current else f"Upgrade to {name}"
        draw_button(d, px+20, py+398, btn_label, w=pw-40, style=btn_style)
    return img


def screen_api_keys():
    img, d = new_canvas()
    draw_navbar(d, "API Keys")
    d.text((30, 68), "API Keys", font=bold_font(24), fill=TEXT)
    d.text((30, 102), "Authenticate programmatic access using X-API-Key header.", font=reg_font(13), fill=MUTED)

    # Create form
    draw_card(d, 30, 128, W-60, 120)
    d.text((50, 142), "Create New API Key", font=bold_font(15), fill=TEXT)
    draw_input(d, 50, 168, 500, "Key name (e.g. Production App)")
    draw_input(d, 566, 168, 340, "Scopes (default: *)")
    draw_button(d, 922, 168, "Generate Key")

    # New key reveal
    rounded_rect(d, 30, 266, W-30, 354, r=14, fill=(40,28,10), outline=YELLOW, width=1)
    d.text((50, 280), "⚠  Save your API key — it will not be shown again", font=bold_font(13), fill=YELLOW)
    key_preview = "mcs_7f3k9pq2xz8wm4nt6rv1ys5aj0bc••••••••••••••••••••••••••••••••••••••"
    rounded_rect(d, 50, 304, W-150, 342, r=10, fill=(20,16,8))
    d.text((66, 316), key_preview, font=reg_font(11), fill=(220,180,80))
    draw_button(d, W-138, 304, "Copy", style="outline")

    # Keys list
    draw_card(d, 30, 370, W-60, 360, title="Your API Keys")
    keys = [
        ("Production App",    "mcs_7f3k••", "*",               "2025-03-18 14:00", "Never"),
        ("Mobile Client",     "mcs_2b9x••", "chat,agents",     "2025-03-15 09:30", "2026-03-15"),
        ("CI/CD Pipeline",    "mcs_5r1m••", "workflows",       "2025-03-10 16:45", "Never"),
        ("Analytics Service", "mcs_9p4w••", "usage,audit",     "2025-02-28 11:00", "2025-06-01"),
    ]
    draw_table_row(d, 46, 414, ["Name","Key Prefix","Scopes","Last Used","Expires",""], [220,130,180,220,140,100], is_header=True)
    for i, (kname, prefix, scopes, last, exp) in enumerate(keys):
        y = 450 + i*56
        draw_table_row(d, 46, y, [kname, prefix, scopes, last, exp, ""], [220,130,180,220,140,100], alt=i%2==0)
        draw_button(d, 46+220+130+180+220+140+8, y+10, "Revoke", style="danger")
    return img


def screen_marketplace():
    img, d = new_canvas()
    draw_navbar(d, "Marketplace")
    d.text((30, 68), "Marketplace", font=bold_font(24), fill=TEXT)

    # Tabs
    tabs = ["Browse","Publish","My Items"]
    tx = 30
    for i, tab in enumerate(tabs):
        active = (i == 0)
        tw_px = int(d.textlength(tab, font=reg_font(13))) + 30
        fill = BRAND if active else PANEL
        brd = BRAND if active else BORDER
        rounded_rect(d, tx, 104, tx+tw_px, 136, r=8, fill=fill, outline=brd, width=1)
        d.text((tx+15, 112), tab, font=bold_font(13) if active else reg_font(13),
               fill=WHITE if active else MUTED)
        tx += tw_px + 8

    # Search + filter
    draw_input(d, 30, 150, 500, "Search marketplace…")
    rounded_rect(d, 546, 150, 726, 192, r=10, fill=PANEL_S, outline=BORDER, width=1)
    d.text((562, 162), "All types  ▾", font=reg_font(12), fill=MUTED)

    items = [
        ("Customer Support Agent", "AGENT",    "AGENT",    "bill_dev",   342, ["support","gpt-4o","multilingual"]),
        ("SEO Blog Generator",     "TEMPLATE", "TEMPLATE", "content_co", 218, ["seo","blogging","marketing"]),
        ("Data Analysis Pipeline", "WORKFLOW", "WORKFLOW", "data_team",  189, ["analytics","reporting"]),
        ("HR Onboarding Bot",      "AGENT",    "AGENT",    "hr_tools",   156, ["hr","onboarding","automation"]),
        ("Email Drafter Pro",      "TEMPLATE", "TEMPLATE", "mktg_ai",    134, ["email","outreach","b2b"]),
        ("Weekly Digest Flow",     "WORKFLOW", "WORKFLOW", "ops_stack",   98, ["automation","digest","slack"]),
    ]
    type_col = {"AGENT": BRAND, "TEMPLATE": PURPLE, "WORKFLOW": YELLOW}

    cols = 3
    card_w = (W - 80) // cols - 10
    for i, (title, itype, badge, author, dl, tags) in enumerate(items):
        row, col = i // cols, i % cols
        cx = 30 + col*(card_w+15)
        cy = 208 + row*220
        rounded_rect(d, cx, cy, cx+card_w, cy+200, r=16, fill=PANEL, outline=BORDER, width=1)
        tc = type_col.get(itype, MUTED)
        draw_badge(d, cx+16, cy+16, itype, color=tc)
        d.text((cx+16+80, cy+18), f"{dl} downloads", font=reg_font(10), fill=MUTED)
        d.text((cx+16, cy+46), title, font=bold_font(14), fill=TEXT)
        d.text((cx+16, cy+70), f"by {author}", font=reg_font(11), fill=MUTED)
        # tags
        tagx = cx+16
        for tag in tags[:3]:
            tw_px = draw_tag(d, tagx, cy+96, tag, color=MUTED)
            tagx += tw_px + 6
        draw_button(d, cx+16, cy+152, "Clone to workspace", style="outline")
    return img


def screen_integrations():
    img, d = new_canvas()
    draw_navbar(d, "Integrations")
    d.text((30, 68), "Integrations", font=bold_font(24), fill=TEXT)
    d.text((30, 100), "Connect to external services and trigger automations.", font=reg_font(13), fill=MUTED)

    # Available tiles
    draw_card(d, 30, 120, W-60, 220, title="Available Integrations")
    integrations = [
        ("💬","Slack",       "1 connected"),
        ("🔗","Webhook",     "2 connected"),
        ("✉️","Email",       "0 connected"),
        ("📱","WhatsApp",    "0 connected"),
        ("📁","Google Drive","0 connected"),
        ("📓","Notion",      "1 connected"),
        ("📚","Confluence",  "0 connected"),
        ("👥","CRM",         "0 connected"),
    ]
    tile_w = (W-100)//8 - 8
    for i, (icon, name, connected) in enumerate(integrations):
        tx = 46 + i*(tile_w+8)
        ty = 168
        brd = BRAND if "connected" in connected and "0" not in connected else BORDER
        rounded_rect(d, tx, ty, tx+tile_w, ty+120, r=12, fill=PANEL_S, outline=brd, width=1)
        d.text((tx+tile_w//2-10, ty+18), icon, font=reg_font(24), fill=WHITE)
        d.text((tx+tile_w//2-20, ty+58), name, font=bold_font(11), fill=TEXT)
        if "0" not in connected:
            d.text((tx+8, ty+86), connected, font=reg_font(9), fill=GREEN)
        else:
            d.text((tx+8, ty+86), "Not connected", font=reg_font(9), fill=MUTED)
        d.text((tx+tile_w//2-14, ty+100), "Configure", font=reg_font(9), fill=BRAND_L)

    # Config form (Slack example)
    draw_card(d, 30, 356, 560, 260, title="Configure: Slack")
    d.text((50, 400), "Integration Name", font=reg_font(12), fill=MUTED)
    draw_input(d, 50, 418, 520, value="Slack Notifications")
    d.text((50, 472), "Webhook URL", font=reg_font(12), fill=MUTED)
    draw_input(d, 50, 490, 520, value="https://hooks.slack.com/services/T0••••/B0••••/••••••")
    draw_button(d, 50, 540, "Connect Integration")

    # Connected list
    draw_card(d, 606, 356, W-636, 340, title="Connected (3)")
    connected_list = [
        ("💬","Slack Notifications",    "SLACK",   True,  "2025-03-18 10:00"),
        ("🔗","Workflow Webhook",        "WEBHOOK", True,  "2025-03-17 16:30"),
        ("📓","Notion Sync",            "NOTION",  False, "2025-03-15 09:00"),
    ]
    for i, (icon, name, itype, enabled, last) in enumerate(connected_list):
        ly = 404 + i*96
        rounded_rect(d, 626, ly, W-50, ly+80, r=10, fill=PANEL_S, outline=BORDER, width=1)
        d.text((642, ly+16), icon, font=reg_font(20), fill=WHITE)
        d.text((678, ly+14), name, font=bold_font(13), fill=TEXT)
        d.text((678, ly+36), f"{itype}  ·  Last: {last}", font=reg_font(10), fill=MUTED)
        draw_toggle(d, W-160, ly+20, on=enabled)
        draw_button(d, W-100, ly+18, "Delete", style="danger")
    return img


def screen_audit():
    img, d = new_canvas()
    draw_navbar(d, "Audit")
    d.text((30, 68), "Audit Logs", font=bold_font(24), fill=TEXT)
    draw_input(d, W-320, 62, 280, "Filter by action…")

    draw_card(d, 30, 108, W-60, 600, title="Activity Log")

    headers = ["Action","Resource","Detail","IP Address","Date"]
    widths  = [220, 160, 370, 140, 200]
    draw_table_row(d, 46, 152, headers, widths, is_header=True)

    logs = [
        ("AGENT_CREATE",       "Agent",    "Created 'Legal Assistant'",          "192.168.1.10", "2025-03-18 14:32"),
        ("CHAT_MESSAGE",       "Session",  "Session: abc123",                     "192.168.1.10", "2025-03-18 14:28"),
        ("DOCUMENT_UPLOAD",    "Document", "product_manual.pdf (2.3 MB)",         "192.168.1.10", "2025-03-18 13:55"),
        ("API_KEY_CREATE",     "ApiKey",   "Created 'Production App'",            "192.168.1.10", "2025-03-18 13:40"),
        ("WORKFLOW_EXECUTE",   "Workflow", "Customer Onboarding Flow",            "192.168.1.10", "2025-03-18 13:22"),
        ("SUBSCRIPTION_UPGRADE","Billing", "FREE → PRO",                         "192.168.1.10", "2025-03-17 16:00"),
        ("ORG_MEMBER_INVITE",  "Org",      "Invited sarah@acme.com as ADMIN",     "192.168.1.10", "2025-03-17 15:30"),
        ("AGENT_DISABLE",      "Agent",    "Disabled 'Data Analyst'",             "192.168.1.10", "2025-03-17 14:10"),
        ("DOCUMENT_DELETE",    "Document", "Deleted old_faq.txt",                 "192.168.1.10", "2025-03-17 13:00"),
        ("USER_LOGIN",         "Auth",     "Successful login",                    "192.168.1.10", "2025-03-17 09:00"),
        ("MARKETPLACE_PUBLISH","Marketplace","Published 'SEO Blog Generator'",    "192.168.1.10", "2025-03-16 18:22"),
        ("INTEGRATION_CREATE", "Integration","Created Slack Notifications",       "192.168.1.10", "2025-03-16 11:45"),
        ("AGENT_UPDATE",       "Agent",    "Updated system prompt",               "192.168.1.10", "2025-03-15 10:30"),
    ]

    action_colors = {
        "AGENT_CREATE": BRAND_L, "CHAT_MESSAGE": MUTED, "DOCUMENT_UPLOAD": PURPLE,
        "API_KEY_CREATE": GREEN, "WORKFLOW_EXECUTE": BRAND_L, "SUBSCRIPTION_UPGRADE": YELLOW,
        "ORG_MEMBER_INVITE": BRAND_L, "AGENT_DISABLE": YELLOW, "DOCUMENT_DELETE": RED,
        "USER_LOGIN": GREEN, "MARKETPLACE_PUBLISH": BRAND_L, "INTEGRATION_CREATE": GREEN,
        "AGENT_UPDATE": BRAND_L,
    }
    for i, (action, res, detail, ip, dt) in enumerate(logs):
        y = 188 + i*34
        if y > 680:
            break
        alt = i % 2 == 0
        bg = (22, 32, 55) if alt else PANEL
        d.rectangle([46, y, 46+sum(widths), y+34], fill=bg)
        d.line([46, y+34, 46+sum(widths), y+34], fill=BORDER, width=1)
        acol = action_colors.get(action, TEXT)
        d.text((58, y+9), action, font=bold_font(10), fill=acol)
        d.text((58+widths[0], y+9), res, font=reg_font(11), fill=TEXT)
        d.text((58+widths[0]+widths[1], y+9), detail[:42], font=reg_font(10), fill=MUTED)
        d.text((58+widths[0]+widths[1]+widths[2], y+9), ip, font=reg_font(10), fill=MUTED)
        d.text((58+widths[0]+widths[1]+widths[2]+widths[3], y+9), dt, font=reg_font(10), fill=MUTED)
    return img


def screen_admin():
    img, d = new_canvas()
    draw_navbar(d, "Admin")
    d.text((30, 68), "Admin Panel", font=bold_font(24), fill=TEXT)

    # Tabs
    tabs = ["Users","Agents","Usage Logs"]
    tx = 30
    for i, tab in enumerate(tabs):
        active = (i == 0)
        tw_px = int(d.textlength(tab, font=reg_font(13))) + 28
        fill = BRAND if active else PANEL_S
        rounded_rect(d, tx, 104, tx+tw_px, 132, r=8, fill=fill, outline=BORDER, width=1)
        d.text((tx+14, 111), tab, font=bold_font(12) if active else reg_font(12),
               fill=WHITE if active else MUTED)
        tx += tw_px + 8

    draw_card(d, 30, 144, W-60, 540, title="All Users")

    # Stats bar
    for i, (val, lbl, col) in enumerate([
        ("47","Total users",BRAND_L), ("38","Active",GREEN), ("9","Inactive",MUTED), ("5","Admins",YELLOW)
    ]):
        sx = 50 + i*230
        draw_stat_card(d, sx, 186, 210, 72, val, lbl, col)

    draw_table_row(d, 46, 274, ["Name","Email","Role","Plan","Agents","Joined"], [200,250,100,100,80,160], is_header=True)
    users = [
        ("Bill Chen",     "bill@acme.com",     "ADMIN", "PRO",        "8",  "2025-01-15"),
        ("Sarah Johnson", "sarah@acme.com",    "USER",  "PRO",        "3",  "2025-01-22"),
        ("Mike Davis",    "mike@acme.com",     "USER",  "FREE",       "2",  "2025-02-01"),
        ("Lisa Wang",     "lisa@acme.com",     "USER",  "ENTERPRISE", "15", "2025-02-10"),
        ("Tom Harris",    "tom@startup.io",    "USER",  "FREE",       "1",  "2025-02-18"),
        ("Amy Brown",     "amy@techcorp.com",  "USER",  "PRO",        "6",  "2025-02-28"),
        ("James Lee",     "james@devco.com",   "USER",  "FREE",       "2",  "2025-03-05"),
        ("Emma Wilson",   "emma@agency.co",    "USER",  "PRO",        "4",  "2025-03-10"),
    ]
    for i, row in enumerate(users):
        y = 310 + i*46
        if y > 660:
            break
        draw_table_row(d, 46, y, list(row), [200,250,100,100,80,160], alt=i%2==0)
    return img


# ─────────────────────────────────────────────────────────────────────────────
# Generate all screen images
# ─────────────────────────────────────────────────────────────────────────────

SCREENS = [
    ("01_login",           screen_login,            "Login Page"),
    ("02_dashboard",       screen_dashboard,         "Dashboard"),
    ("03_agent_detail",    screen_agent_detail,      "Agent Detail"),
    ("04_agent_settings",  screen_agent_settings,    "Agent Settings"),
    ("05_knowledge_bases", screen_knowledge_bases,   "Knowledge Bases"),
    ("06_workflows",       screen_workflows,          "Workflows"),
    ("07_workflow_builder",screen_workflow_builder,   "Workflow Builder"),
    ("08_scheduled_jobs",  screen_scheduled_jobs,    "Scheduled Jobs"),
    ("09_templates",       screen_templates,          "Prompt Templates"),
    ("10_usage",           screen_usage,              "Usage & Cost"),
    ("11_organization",    screen_organization,       "Organizations"),
    ("12_billing",         screen_billing,            "Billing"),
    ("13_api_keys",        screen_api_keys,           "API Keys"),
    ("14_marketplace",     screen_marketplace,        "Marketplace"),
    ("15_integrations",    screen_integrations,       "Integrations"),
    ("16_audit",           screen_audit,              "Audit Logs"),
    ("17_admin",           screen_admin,              "Admin Panel"),
]

def generate_all_images():
    paths = {}
    for slug, fn, label in SCREENS:
        img = fn()
        path = IMG_DIR / f"{slug}.png"
        img.save(path, "PNG", optimize=True)
        paths[slug] = (str(path), label)
        print(f"  [img] {slug}.png")
    return paths

# ─────────────────────────────────────────────────────────────────────────────
# User guide content
# ─────────────────────────────────────────────────────────────────────────────

GUIDE_SECTIONS = [
    {
        "title": "1. Getting Started",
        "slug": "01_login",
        "caption": "Figure 1 — Login page",
        "intro": (
            "Welcome to MCS AI Platform — a complete SaaS solution for building, deploying, "
            "and managing AI-powered agents. This guide walks you through every feature of "
            "the platform with step-by-step instructions."
        ),
        "subsections": [
            {
                "heading": "1.1 Creating an Account",
                "body": (
                    "Navigate to the platform URL and click Register on the login page. "
                    "Enter your full name, email address, and a secure password. "
                    "Once registered, you are automatically logged in."
                ),
            },
            {
                "heading": "1.2 Logging In",
                "body": (
                    "Enter your email and password on the Sign in page and click the Sign in button. "
                    "Your session token is stored locally and persists until you click Log out "
                    "in the navigation bar."
                ),
            },
        ],
    },
    {
        "title": "2. Dashboard",
        "slug": "02_dashboard",
        "caption": "Figure 2 — Dashboard showing agent cards",
        "intro": (
            "The Dashboard is your home screen. It displays all your AI agents as cards "
            "with quick access to chat and settings. Disabled agents appear greyed out "
            "with a 'Disabled' badge."
        ),
        "subsections": [
            {
                "heading": "2.1 Creating an Agent",
                "body": (
                    "Click the + New Agent button (top right). Enter a name and optionally a "
                    "description. The agent is created with sensible defaults: gpt-3.5-turbo model, "
                    "temperature 0.7, memory enabled."
                ),
            },
            {
                "heading": "2.2 Agent Status Indicators",
                "body": (
                    "A green dot next to the agent name indicates it is active and ready to chat. "
                    "A grey dot and 'Disabled' badge mean the agent has been disabled via the "
                    "Settings → Danger Zone section. Disabled agents cannot be used for chat."
                ),
            },
        ],
    },
    {
        "title": "3. Agent Detail & Chat",
        "slug": "03_agent_detail",
        "caption": "Figure 3 — Agent detail: documents on the left, chat preview on the right",
        "intro": (
            "Click an agent card from the Dashboard to open its detail page. Here you can "
            "manage the agent's knowledge documents and start or review chat conversations."
        ),
        "subsections": [
            {
                "heading": "3.1 Uploading Documents",
                "body": (
                    "The document upload follows a three-step process: (1) Click Upload Document "
                    "to open the file picker, (2) select a PDF, TXT, or Markdown file (max 20 MB), "
                    "(3) click Add to confirm. The document is parsed, split into chunks, and "
                    "embedded using the text-embedding-3-small model for semantic search."
                ),
            },
            {
                "heading": "3.2 Downloading & Deleting Documents",
                "body": (
                    "Each document row in the table has Download (saves the original file to your "
                    "computer) and Delete (removes the file, all chunks, and the database record) "
                    "buttons. Deletion is permanent."
                ),
            },
            {
                "heading": "3.3 Chatting with an Agent",
                "body": (
                    "Click Start Chat (top right) to open a chat session. Type your message in the "
                    "input bar and press Send. The agent responds using its configured LLM, system "
                    "prompt, knowledge base, and enabled tools. The last 10 messages are kept as "
                    "context; sessions exceeding 20 messages are automatically summarised."
                ),
            },
        ],
    },
    {
        "title": "4. Agent Settings",
        "slug": "04_agent_settings",
        "caption": "Figure 4 — Agent settings: model configuration, tools, and knowledge base bindings",
        "intro": (
            "Click Settings on any agent card to open the full configuration page. "
            "All changes must be saved by clicking Save Changes."
        ),
        "subsections": [
            {
                "heading": "4.1 Model Configuration",
                "body": (
                    "Agent Name and Description are used for display purposes. The System Prompt "
                    "defines the agent's persona and behaviour. Model selects the LLM (e.g. gpt-4o, "
                    "gpt-4o-mini, gpt-3.5-turbo). Temperature (0.0–2.0) controls response "
                    "creativity — lower values are more deterministic."
                ),
            },
            {
                "heading": "4.2 Memory & Tools",
                "body": (
                    "Memory Enabled — when toggled on, the platform injects recent conversation "
                    "history into each LLM call for contextual continuity. "
                    "Tools Enabled — allows the agent to use registered tool executors "
                    "(Knowledge Search, Web Search, Calculator). Individual tools can be toggled "
                    "on or off per agent."
                ),
            },
            {
                "heading": "4.3 Knowledge Base Bindings",
                "body": (
                    "Link one or more Knowledge Bases to this agent. When the Knowledge Search "
                    "tool runs, it searches across all linked KBs using cosine similarity on "
                    "the embedded document chunks."
                ),
            },
            {
                "heading": "4.4 Danger Zone",
                "body": (
                    "Disable Agent — makes the agent unavailable for chat without deleting it. "
                    "Delete Agent — permanently removes the agent and all associated data. "
                    "Both actions require confirmation."
                ),
            },
        ],
    },
    {
        "title": "5. Knowledge Bases",
        "slug": "05_knowledge_bases",
        "caption": "Figure 5 — Knowledge bases management",
        "intro": (
            "Knowledge Bases are collections of documents that can be shared across multiple agents. "
            "Navigate to Knowledge Bases in the top navigation bar."
        ),
        "subsections": [
            {
                "heading": "5.1 Creating a Knowledge Base",
                "body": (
                    "Enter a name (required) and optional description in the Create form at the top, "
                    "then click Create. The KB starts empty — add documents by uploading them "
                    "to an agent that has this KB linked."
                ),
            },
            {
                "heading": "5.2 Linking to Agents",
                "body": (
                    "Go to Agent Settings → Knowledge Bases section and toggle on the KBs you want "
                    "the agent to search. An agent can be linked to multiple KBs simultaneously."
                ),
            },
        ],
    },
    {
        "title": "6. Workflows",
        "slug": "06_workflows",
        "caption": "Figure 6 — Workflow list",
        "intro": (
            "Workflows let you chain multiple AI steps into an automated pipeline. "
            "Each step can call the LLM, search knowledge, fetch web results, "
            "summarise content, or make HTTP requests."
        ),
        "subsections": [
            {
                "heading": "6.1 Creating a Workflow",
                "body": (
                    "Click Create on the Workflows page. Enter a name and optional description. "
                    "The workflow opens with zero steps — click + Add step to start building."
                ),
            },
        ],
    },
    {
        "title": "7. Workflow Builder",
        "slug": "07_workflow_builder",
        "caption": "Figure 7 — Workflow builder with step editor and execution panel",
        "intro": (
            "The Workflow Builder allows you to visually construct multi-step AI pipelines "
            "and execute them with a single click."
        ),
        "subsections": [
            {
                "heading": "7.1 Adding Steps",
                "body": (
                    "Click + Add step. For each step configure: Name (descriptive label), "
                    "Type (KNOWLEDGE_SEARCH, WEB_SEARCH, LLM_PROMPT, SUMMARIZE, HTTP_REQUEST), "
                    "Input Template (use {{variableName}} to reference previous step outputs or "
                    "the initial {{input}}), Output Key (name under which this step's result "
                    "is stored for downstream steps)."
                ),
            },
            {
                "heading": "7.2 Executing a Workflow",
                "body": (
                    "Type your initial input in the Execute panel, then click Run Workflow. "
                    "Execution is asynchronous — the status updates every 2 seconds automatically. "
                    "The Execution History section shows step-by-step results, durations, and "
                    "any errors, as well as the final output."
                ),
            },
            {
                "heading": "7.3 Step Variable Syntax",
                "body": (
                    "Use {{input}} to reference the workflow's initial input. "
                    "Use {{outputKey}} to reference the output of a previous step, "
                    "where outputKey is the Output Key configured on that step. "
                    "Example: if step 1 has Output Key search_result, step 2 can use "
                    "{{search_result}} in its Input Template."
                ),
            },
        ],
    },
    {
        "title": "8. Scheduled Jobs",
        "slug": "08_scheduled_jobs",
        "caption": "Figure 8 — Scheduled jobs with HOURLY/DAILY/WEEKLY options",
        "intro": (
            "Scheduled Jobs automatically trigger workflows on a recurring schedule. "
            "The scheduler checks for due jobs every 60 seconds."
        ),
        "subsections": [
            {
                "heading": "8.1 Creating a Scheduled Job",
                "body": (
                    "Select the workflow to run, choose a schedule type (HOURLY, DAILY, WEEKLY), "
                    "give the job a name, and click Schedule. The first run is queued immediately; "
                    "subsequent runs are calculated from the last run time."
                ),
            },
            {
                "heading": "8.2 Managing Jobs",
                "body": (
                    "Each job card shows the last run time and next scheduled run. "
                    "Click Disable to pause a job without deleting it. Click Enable to resume. "
                    "Click Delete to permanently remove the job."
                ),
            },
        ],
    },
    {
        "title": "9. Prompt Templates",
        "slug": "09_templates",
        "caption": "Figure 9 — Prompt template editor with variable detection and render panel",
        "intro": (
            "Prompt Templates let you define reusable prompt blueprints with dynamic "
            "placeholders. Use them in workflows or render them directly in the UI."
        ),
        "subsections": [
            {
                "heading": "9.1 Creating a Template",
                "body": (
                    "Write your template content using {{variableName}} placeholders — "
                    "for example: Write a {{tone}} email about {{topic}} for {{audience}}. "
                    "The platform automatically detects and displays all variable names as badges. "
                    "Click Create to save."
                ),
            },
            {
                "heading": "9.2 Rendering a Template",
                "body": (
                    "Click a template in the list to select it. The Render panel appears on the "
                    "right with an input field for each variable. Fill in the values and click "
                    "Render Template to see the final output immediately."
                ),
            },
        ],
    },
    {
        "title": "10. Usage & Cost",
        "slug": "10_usage",
        "caption": "Figure 10 — Usage dashboard with stats, model charts, and log table",
        "intro": (
            "The Usage & Cost page gives you complete visibility into platform consumption "
            "and estimated spending across all agents and models."
        ),
        "subsections": [
            {
                "heading": "10.1 Summary Cards",
                "body": (
                    "The three stat cards at the top show total API requests made, total tokens "
                    "consumed, and estimated cost in USD since account creation."
                ),
            },
            {
                "heading": "10.2 Breakdown Charts",
                "body": (
                    "Tokens by Model shows the proportional share of each LLM model's token usage "
                    "with a progress bar. Tokens by Agent shows the same breakdown per agent. "
                    "Both charts include cost figures."
                ),
            },
            {
                "heading": "10.3 Recent Logs",
                "body": (
                    "The log table shows the 50 most recent API calls with agent, model, token count, "
                    "estimated cost, and timestamp. Use this for detailed cost attribution."
                ),
            },
        ],
    },
    {
        "title": "11. Organizations",
        "slug": "11_organization",
        "caption": "Figure 11 — Organization management with member roles",
        "intro": (
            "Organizations enable team collaboration by grouping users under a shared workspace "
            "with role-based access control."
        ),
        "subsections": [
            {
                "heading": "11.1 Creating an Organization",
                "body": (
                    "Click Create Organization, enter a name and optional description. "
                    "You become the OWNER automatically and are the only member initially."
                ),
            },
            {
                "heading": "11.2 Inviting Members",
                "body": (
                    "Enter the email address of a registered platform user and select their role: "
                    "MEMBER (read-only) or ADMIN (can invite/remove members). "
                    "The user must already have an account on the platform."
                ),
            },
            {
                "heading": "11.3 Roles",
                "body": (
                    "OWNER — full control, cannot be removed. "
                    "ADMIN — can invite and remove members, manage org settings. "
                    "MEMBER — can access shared resources within the organisation."
                ),
            },
        ],
    },
    {
        "title": "12. Billing & Subscription",
        "slug": "12_billing",
        "caption": "Figure 12 — Billing page with plan comparison cards",
        "intro": (
            "The Billing page shows your current subscription plan and allows you to upgrade "
            "to unlock higher limits."
        ),
        "subsections": [
            {
                "heading": "12.1 Plan Comparison",
                "body": (
                    "FREE ($0/mo): 5 agents, 100 requests/day, 50 documents. "
                    "PRO ($29/mo): 50 agents, 5,000 requests/day, 1,000 documents, API access. "
                    "ENTERPRISE ($149/mo): 500 agents, 100,000 requests/day, 10,000 documents, "
                    "SSO/SAML, SLA guarantee."
                ),
            },
            {
                "heading": "12.2 Upgrading",
                "body": (
                    "Click the Upgrade button on any plan card and confirm when prompted. "
                    "The plan change takes effect immediately. Billing integration with Stripe "
                    "is configured by the platform administrator."
                ),
            },
        ],
    },
    {
        "title": "13. API Keys",
        "slug": "13_api_keys",
        "caption": "Figure 13 — API key management with one-time reveal panel",
        "intro": (
            "API Keys let you authenticate programmatic access to the platform without "
            "using your password. Pass the key as the X-API-Key request header."
        ),
        "subsections": [
            {
                "heading": "13.1 Generating an API Key",
                "body": (
                    "Enter a descriptive name and optional scope list (comma-separated, default *), "
                    "then click Generate Key. The full key is displayed once in a yellow panel — "
                    "copy it immediately as it cannot be retrieved again."
                ),
            },
            {
                "heading": "13.2 Using an API Key",
                "body": (
                    "Include the header X-API-Key: mcs_<your_key> in every request. "
                    "API keys support the same endpoints as JWT authentication. "
                    "The key's Last Used timestamp is updated on each use."
                ),
            },
            {
                "heading": "13.3 Revoking Keys",
                "body": (
                    "Click Revoke on any key to permanently delete it. "
                    "Any application using that key will immediately lose access. "
                    "Revocation cannot be undone."
                ),
            },
        ],
    },
    {
        "title": "14. Marketplace",
        "slug": "14_marketplace",
        "caption": "Figure 14 — Marketplace with agent, template, and workflow listings",
        "intro": (
            "The Marketplace is a public library of community-published agents, prompt templates, "
            "and workflows. Browse without logging in; publish and clone with an account."
        ),
        "subsections": [
            {
                "heading": "14.1 Browsing & Cloning",
                "body": (
                    "Use the search box to find items by title or the type filter to narrow "
                    "by AGENT, TEMPLATE, or WORKFLOW. Click Clone to workspace to copy an item's "
                    "configuration into your account. The clone counter increments on each download."
                ),
            },
            {
                "heading": "14.2 Publishing",
                "body": (
                    "Switch to the Publish tab. Select the type, enter a title, description, "
                    "category, and comma-separated tags, then click Publish. "
                    "Your item appears in the public listing immediately."
                ),
            },
            {
                "heading": "14.3 Managing Your Items",
                "body": (
                    "The My Items tab lists everything you have published. "
                    "Click Remove to unpublish an item — it disappears from the marketplace "
                    "but is not deleted from your account."
                ),
            },
        ],
    },
    {
        "title": "15. Integrations",
        "slug": "15_integrations",
        "caption": "Figure 15 — Integrations page with tile grid and Slack configuration form",
        "intro": (
            "Integrations connect the platform to external services. Configured integrations "
            "are triggered automatically when workflows complete, sending output to Slack, "
            "webhooks, email, and more."
        ),
        "subsections": [
            {
                "heading": "15.1 Available Integration Types",
                "body": (
                    "Slack — post messages via Incoming Webhooks. "
                    "Webhook (generic) — POST JSON to any URL when a workflow completes. "
                    "Email — send results via SMTP. "
                    "WhatsApp — send via Meta Cloud API. "
                    "Google Drive, Notion, Confluence — connect your document platforms. "
                    "CRM — Salesforce, HubSpot via API key."
                ),
            },
            {
                "heading": "15.2 Adding an Integration",
                "body": (
                    "Click any integration tile to open its configuration form. "
                    "Enter the integration name and fill in the required fields "
                    "(webhook URL, API token, etc.). Click Connect to save. "
                    "Sensitive fields are masked with password input styling."
                ),
            },
            {
                "heading": "15.3 Enabling & Disabling",
                "body": (
                    "Each connected integration has an Enable/Disable toggle. "
                    "Disabled integrations are not triggered even when workflows complete. "
                    "This is useful for temporarily pausing an integration without deleting it."
                ),
            },
        ],
    },
    {
        "title": "16. Audit Logs",
        "slug": "16_audit",
        "caption": "Figure 16 — Audit log table with colour-coded actions",
        "intro": (
            "Audit Logs provide a complete, tamper-visible record of every action taken "
            "in your account, including the IP address and timestamp of each event."
        ),
        "subsections": [
            {
                "heading": "16.1 Reading the Audit Log",
                "body": (
                    "Each row shows the action type (colour-coded by category), the resource "
                    "it affected, a human-readable detail string, the originating IP address, "
                    "and the exact date/time."
                ),
            },
            {
                "heading": "16.2 Filtering",
                "body": (
                    "Type in the Filter by action box (top right) to narrow the log to "
                    "specific action types such as AGENT, DOCUMENT, or WORKFLOW."
                ),
            },
        ],
    },
    {
        "title": "17. Admin Panel",
        "slug": "17_admin",
        "caption": "Figure 17 — Admin panel showing the Users tab",
        "intro": (
            "The Admin Panel is accessible to all authenticated users and provides "
            "platform-wide visibility across users, agents, and usage logs."
        ),
        "subsections": [
            {
                "heading": "17.1 Users Tab",
                "body": (
                    "Lists all registered users with their email, role, subscription plan, "
                    "number of agents, and registration date. Summary stats at the top "
                    "show total, active, inactive, and admin user counts."
                ),
            },
            {
                "heading": "17.2 Agents Tab",
                "body": (
                    "Shows all agents across all users, useful for monitoring platform-wide "
                    "agent usage and identifying inactive or misconfigured agents."
                ),
            },
            {
                "heading": "17.3 Usage Logs Tab",
                "body": (
                    "Displays all API calls across all users — valuable for cost allocation, "
                    "quota monitoring, and identifying heavy consumers."
                ),
            },
        ],
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# Word document generator
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def generate_word_guide(image_paths, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(PROJECT)
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("User Guide")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x37, 0x63, 0xEB)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(f"Version {VERSION}  ·  {TODAY}").font.size = Pt(11)

    doc.add_page_break()

    # ── Table of Contents ────────────────────────────────────────────────────
    toc_h = doc.add_heading("Table of Contents", level=1)
    toc_h.runs[0].font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)
    for gs in GUIDE_SECTIONS:
        p = doc.add_paragraph(gs["title"], style='List Number')
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── Sections ─────────────────────────────────────────────────────────────
    for gs in GUIDE_SECTIONS:
        h = doc.add_heading(gs["title"], level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)

        # Intro paragraph
        intro_p = doc.add_paragraph(gs["intro"])
        intro_p.paragraph_format.space_after = Pt(8)

        # Screenshot
        slug = gs["slug"]
        if slug in image_paths:
            img_path, _ = image_paths[slug]
            try:
                doc.add_picture(img_path, width=Inches(6.2))
                cap_p = doc.add_paragraph(gs["caption"])
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.runs[0].font.size = Pt(9)
                cap_p.runs[0].font.italic = True
                cap_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"[Image: {img_path}]")

        # Subsections
        for sub in gs.get("subsections", []):
            sh = doc.add_heading(sub["heading"], level=2)
            sh.runs[0].font.color.rgb = RGBColor(0x34, 0x49, 0x6B)
            body_p = doc.add_paragraph(sub["body"])
            body_p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

    doc.save(output_path)
    print(f"Word guide saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF guide generator
# ─────────────────────────────────────────────────────────────────────────────

NAVY  = colors.HexColor('#1A2B4A')
STEEL = colors.HexColor('#34496B')
BRAND_C = colors.HexColor('#2563EB')
GRAY  = colors.HexColor('#64748B')
MID   = colors.HexColor('#E2E8F0')

def get_guide_pdf_styles():
    styles = getSampleStyleSheet()
    base = styles['Normal']
    s = {}
    s['Title']    = ParagraphStyle('GT', parent=base, fontSize=32, leading=42,
                                    textColor=NAVY, fontName='Helvetica-Bold',
                                    spaceAfter=6, alignment=TA_CENTER)
    s['Subtitle'] = ParagraphStyle('GS', parent=base, fontSize=20, leading=28,
                                    textColor=BRAND_C, fontName='Helvetica',
                                    spaceAfter=4, alignment=TA_CENTER)
    s['Meta']     = ParagraphStyle('GM', parent=base, fontSize=11, textColor=GRAY,
                                    fontName='Helvetica', spaceAfter=30, alignment=TA_CENTER)
    s['H1']       = ParagraphStyle('GH1', parent=base, fontSize=17, leading=24,
                                    textColor=NAVY, fontName='Helvetica-Bold',
                                    spaceBefore=22, spaceAfter=8)
    s['H2']       = ParagraphStyle('GH2', parent=base, fontSize=12, leading=18,
                                    textColor=STEEL, fontName='Helvetica-Bold',
                                    spaceBefore=14, spaceAfter=5)
    s['Body']     = ParagraphStyle('GB', parent=base, fontSize=10, leading=15,
                                    textColor=colors.HexColor('#1E293B'),
                                    fontName='Helvetica', spaceAfter=8,
                                    alignment=TA_JUSTIFY)
    s['Caption']  = ParagraphStyle('GC', parent=base, fontSize=8.5, leading=13,
                                    textColor=GRAY, fontName='Helvetica-Oblique',
                                    spaceAfter=10, alignment=TA_CENTER)
    s['TOCTitle'] = ParagraphStyle('GTOC', parent=base, fontSize=15, textColor=NAVY,
                                    fontName='Helvetica-Bold', spaceAfter=10)
    s['TOCItem']  = ParagraphStyle('GTOCI', parent=base, fontSize=10, textColor=STEEL,
                                    fontName='Helvetica', leftIndent=14, spaceAfter=3)
    return s


def generate_pdf_guide(image_paths, output_path):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        title=f"{PROJECT} User Guide",
        author="MCS Engineering",
    )

    styles = get_guide_pdf_styles()
    story = []

    # Cover
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(PROJECT, styles['Title']))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph("User Guide", styles['Subtitle']))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph(f"Version {VERSION}  ·  {TODAY}", styles['Meta']))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BRAND_C, spaceAfter=20))

    # TOC
    story.append(Paragraph("Table of Contents", styles['TOCTitle']))
    for gs in GUIDE_SECTIONS:
        story.append(Paragraph(gs["title"], styles['TOCItem']))
    story.append(PageBreak())

    # Sections
    page_width = A4[0] - 5*cm

    for gs in GUIDE_SECTIONS:
        story.append(Paragraph(gs["title"], styles['H1']))
        story.append(HRFlowable(width="100%", thickness=0.5, color=MID, spaceAfter=8))
        story.append(Paragraph(gs["intro"], styles['Body']))

        # Image
        slug = gs["slug"]
        if slug in image_paths:
            img_path, _ = image_paths[slug]
            try:
                story.append(RLImage(img_path, width=page_width, height=page_width * (H/W)))
                story.append(Paragraph(gs["caption"], styles['Caption']))
                story.append(Spacer(1, 0.3*cm))
            except Exception as e:
                story.append(Paragraph(f"[Image unavailable: {e}]", styles['Caption']))

        # Subsections
        for sub in gs.get("subsections", []):
            story.append(Paragraph(sub["heading"], styles['H2']))
            story.append(Paragraph(sub["body"], styles['Body']))

        story.append(Spacer(1, 0.5*cm))

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(GRAY)
        canvas.drawString(2.5*cm, 1.5*cm, f"{PROJECT} User Guide — v{VERSION}")
        canvas.drawRightString(A4[0] - 2.5*cm, 1.5*cm, f"Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    print(f"PDF guide saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("Generating UI mockup images…")
    image_paths = generate_all_images()

    word_out = str(BASE / "MCS_AI_Platform_User_Guide.docx")
    pdf_out  = str(BASE / "MCS_AI_Platform_User_Guide.pdf")

    print("\nBuilding Word user guide…")
    generate_word_guide(image_paths, word_out)

    print("\nBuilding PDF user guide…")
    generate_pdf_guide(image_paths, pdf_out)

    print(f"\nDone! Files generated:\n  {word_out}\n  {pdf_out}")
